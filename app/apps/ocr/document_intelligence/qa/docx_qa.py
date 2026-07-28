"""
Post-generation DOCX QA gate.

Checks a rendered .docx against the DocumentAST it was built from, per the
"Quality Assurance" section of the Semantic DOCX redesign plan. This is a
structural/XML-level check (no LibreOffice/Word needed) meant to run in CI
as a fail-hard gate on the semantic renderer (docx.py) -- the whole point
of the redesign is that normal text must never end up in a Text Box/Shape
or at an absolute position, and this is what actually enforces that
instead of relying on someone noticing in a manual review.

mode="visual" (the absolute-layout renderer, docx_absolute.py)
intentionally uses text boxes and absolute positioning, so those two
checks are reported as skipped -- not failed -- for that mode.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Literal

from docx import Document
from docx.oxml.ns import qn

from ..ast import DocumentAST
from ..latex_omml import LatexConversionError, latex_to_omml
from ..layout import LayoutType
from ..renderers.docx import _collect_header_footer
from ..structure.header_footer import detect_header_footer_regions

_PERSIAN_RE = re.compile(r"[؀-ۿݐ-ݿ]")
_OMATH_TAG_RE = re.compile(r"<m:oMath[ >]")
_INSTR_PAGE_RE = re.compile(r"<w:instrText[^>]*>\s*PAGE\s*</w:instrText>")
_MARKDOWN_CHARS_RE = re.compile(r"[*`]")
# Inline $latex$/$$latex$$ spans become a real m:oMath object, not
# literal text -- stripped entirely (not just the $ delimiters) before
# comparing AST text against rendered body text, since <m:t> math glyphs
# live outside the <w:t> elements _body_reading_order collects.
_MATH_SPAN_RE = re.compile(r"\${1,2}[^$\n]+?\${1,2}")
_WHITESPACE_RE = re.compile(r"\s+")

# Node types whose .text is expected to land somewhere in the document body
# as ordinary flowing text (as opposed to header/footer/page_number, which
# are conditionally promoted/skipped, or table/figure, which are checked
# separately by their own structural checks).
_TEXT_NODE_TYPES = {
    LayoutType.title,
    LayoutType.heading,
    LayoutType.paragraph,
    LayoutType.reference,
    LayoutType.code,
    LayoutType.table_caption,
    LayoutType.table_footnote,
    LayoutType.figure_caption,
    LayoutType.unknown,
}


@dataclass
class QACheck:
    """Represent QACheck."""

    name: str
    passed: bool
    detail: str = ""


@dataclass
class DocxQAReport:
    """Represent DocxQAReport."""

    mode: str
    checks: list[QACheck] = field(default_factory=list)

    @property
    def passed(self) -> bool:
        """Perform passed."""
        return all(c.passed for c in self.checks)

    def failures(self) -> list[QACheck]:
        """Perform failures."""
        return [c for c in self.checks if not c.passed]

    def to_dict(self) -> dict:
        """Perform to dict."""
        return {
            "mode": self.mode,
            "passed": self.passed,
            "checks": [
                {"name": c.name, "passed": c.passed, "detail": c.detail}
                for c in self.checks
            ],
        }


def run_docx_qa(
    ast: DocumentAST,
    docx_bytes: bytes,
    mode: Literal["semantic", "visual"] = "semantic",
) -> DocxQAReport:
    """Run the full QA checklist and return a report (never raises on failed checks)."""
    doc = Document(BytesIO(docx_bytes))
    xml = doc.element.xml
    body_items = _body_reading_order(doc)

    checks = [
        _check_no_text_boxes(xml, mode),
        _check_no_absolute_positioning(xml, mode),
        _check_reading_order(ast, body_items),
        _check_all_text_consumed(ast, doc, body_items),
        _check_tables_are_real(ast, doc),
        _check_images_are_real_pictures(ast, doc, mode),
        _check_formulas_are_real_omml(ast, xml),
        _check_rtl_bidi(doc),
        _check_header_footer_promotion(ast, doc, body_items),
        _check_page_number_is_real_field(ast, doc, xml, mode),
    ]
    return DocxQAReport(mode=mode, checks=checks)


def _normalize(text: str) -> str:
    # Empty-string, not a space: the real OMML replacement is zero-width in
    # the <w:t>-based text walk (an inline m:oMath contributes no <w:t> at
    # all), so "($y_i = 1$)" renders as literal "()" with nothing between
    # the parens, not "( )" with a space.
    text = _MATH_SPAN_RE.sub("", text or "")
    text = _MARKDOWN_CHARS_RE.sub("", text)
    return _WHITESPACE_RE.sub(" ", text).strip()


def _body_reading_order(doc: Document) -> list[dict]:
    """
    Direct-children walk of doc.element.body, in true document order.

    Unlike doc.paragraphs/doc.tables, which each only walk their own kind
    and lose the interleaved order between them.
    """
    items: list[dict] = []
    for child in doc.element.body:
        if child.tag == qn("w:p"):
            item = {"kind": "paragraph", "text": _element_text(child), "xml": child}
            items.append(item)
        elif child.tag == qn("w:tbl"):
            cells = [_element_text(tc) for tc in child.iter(qn("w:tc"))]
            items.append({"kind": "table", "text": " | ".join(cells), "xml": child})
    return items


def _element_text(el: object) -> str:
    r"""
    Reconstruct visible text the way python-docx's Paragraph.text property does.

    <w:t> contributes its text, <w:tab/> a tab, and <w:br/>/<w:cr/> a
    newline. A raw <w:t>-only walk (the previous version of this
    function) silently drops line breaks -- e.g. one AST node's text
    containing an embedded "\\n" becomes a real <w:br/> when rendered
    (python-docx's run-text setter does this automatically), and
    comparing that against the un-reconstructed rendered text would
    wrongly report the block as missing even though it round-trips
    correctly.
    """
    parts = []
    for node in el.iter():
        tag = node.tag
        if tag == qn("w:t"):
            parts.append(node.text or "")
        elif tag in (qn("w:br"), qn("w:cr")):
            parts.append("\n")
        elif tag == qn("w:tab"):
            parts.append("\t")
    return "".join(parts)


def _expected_text_nodes(ast: DocumentAST) -> list[str]:
    texts: list[str] = []
    for page in ast.pages:
        for node in page.nodes:
            if node.type == LayoutType.list:
                texts.extend(c.text.strip() for c in node.children if c.text.strip())
            elif node.type in _TEXT_NODE_TYPES and node.text.strip():
                texts.append(node.text.strip())
    return texts


def _check_no_text_boxes(xml: str, mode: str) -> QACheck:
    if mode != "semantic":
        return QACheck("no_text_boxes", True, "skipped (visual mode uses text boxes)")
    count = xml.count("txbxContent")
    detail = "" if count == 0 else f"{count} txbxContent occurrence(s) found"
    return QACheck("no_text_boxes", count == 0, detail)


def _check_no_absolute_positioning(xml: str, mode: str) -> QACheck:
    if mode != "semantic":
        return QACheck(
            "no_absolute_positioning",
            True,
            "skipped (visual mode uses absolute positioning by design)",
        )
    found = "position:absolute" in xml
    detail = "found 'position:absolute' in output XML" if found else ""
    return QACheck("no_absolute_positioning", not found, detail)


def _check_reading_order(ast: DocumentAST, body_items: list[dict]) -> QACheck:
    expected = [_normalize(t) for t in _expected_text_nodes(ast)]
    actual_texts = [
        _normalize(item["text"]) for item in body_items if item["kind"] == "paragraph"
    ]

    idx = 0
    for text in actual_texts:
        if idx < len(expected) and expected[idx] and expected[idx] in text:
            idx += 1
    passed = idx == len(expected)
    detail = "" if passed else f"matched {idx}/{len(expected)} expected blocks in order"
    return QACheck("reading_order_preserved", passed, detail)


def _check_all_text_consumed(
    ast: DocumentAST, doc: Document, body_items: list[dict]
) -> QACheck:
    expected = _expected_text_nodes(ast)
    haystack = _normalize(" \n ".join(item["text"] for item in body_items))
    header = doc.sections[0].header
    footer = doc.sections[0].footer
    if header.paragraphs:
        haystack += " \n " + _normalize(header.paragraphs[0].text)
    if footer.paragraphs:
        haystack += " \n " + _normalize(footer.paragraphs[0].text)

    missing = [t for t in expected if _normalize(t) not in haystack]
    passed = not missing
    detail = "" if passed else f"{len(missing)} block(s) missing: {missing[:3]!r}"
    return QACheck("all_ast_text_consumed", passed, detail)


def _check_tables_are_real(ast: DocumentAST, doc: Document) -> QACheck:
    expected = sum(
        1
        for page in ast.pages
        for node in page.nodes
        if node.type == LayoutType.table and node.rows
    )
    actual = len(doc.tables)
    passed = actual >= expected
    detail = "" if passed else f"expected >= {expected} real tables, found {actual}"
    return QACheck("tables_are_real_word_tables", passed, detail)


def _check_images_are_real_pictures(
    ast: DocumentAST, doc: Document, mode: str
) -> QACheck:
    expected = sum(
        1
        for page in ast.pages
        for node in page.nodes
        if node.type in (LayoutType.figure, LayoutType.chart)
        and node.asset_path
        and Path(node.asset_path).exists()
    )
    if mode == "semantic":
        actual = len(doc.inline_shapes)
        detail_kind = "inline pictures"
    else:
        actual = doc.element.xml.count("v:imagedata")
        detail_kind = "embedded images"
    passed = actual >= expected
    detail = "" if passed else f"expected >= {expected} {detail_kind}, found {actual}"
    return QACheck("images_are_real_pictures", passed, detail)


def _check_formulas_are_real_omml(ast: DocumentAST, xml: str) -> QACheck:
    expected_success = 0
    for page in ast.pages:
        for node in page.nodes:
            if node.type == LayoutType.formula and node.latex:
                try:
                    latex_to_omml(node.latex)
                    expected_success += 1
                except LatexConversionError:
                    pass  # falls back to styled text by design — not counted here
    actual = len(_OMATH_TAG_RE.findall(xml))
    passed = actual >= expected_success
    detail = "" if passed else f"expected >={expected_success} OMML, found {actual}"
    return QACheck("formulas_are_real_omml", passed, detail)


def _paragraph_own_text(p: object) -> str:
    """
    Text from this <w:p>'s own direct <w:r> runs only.

    Not text belonging to a further-nested <w:p> hosted inside one of its
    runs' <w:pict> (relevant only for the visual/absolute-layout
    renderer, where the outer wrapper paragraph has no text of its own --
    only the nested txbxContent paragraph does, with its own independent
    pPr).
    """
    return "".join(
        t.text or "" for r in p.findall(qn("w:r")) for t in r.findall(qn("w:t"))
    )


def _check_rtl_bidi(doc: Document) -> QACheck:
    """
    Every real <w:p> containing Persian text must carry <w:bidi/> on its own pPr.

    Checked over all <w:p> elements in the body, recursively, so both the
    flat semantic renderer and the nested-textbox visual renderer are
    covered by the same check.
    """
    offenders = []
    for p in doc.element.body.iter(qn("w:p")):
        text = _paragraph_own_text(p)
        if not _PERSIAN_RE.search(text):
            continue
        ppr = p.find(qn("w:pPr"))
        has_bidi = ppr is not None and ppr.find(qn("w:bidi")) is not None
        if not has_bidi:
            offenders.append(text[:30])
    passed = not offenders
    detail = "" if passed else f"{len(offenders)} missing w:bidi: {offenders[:3]!r}"
    return QACheck("rtl_bidi_correct", passed, detail)


def _check_header_footer_promotion(
    ast: DocumentAST, doc: Document, body_items: list[dict]
) -> QACheck:
    header_text, footer_text = _collect_header_footer(ast)
    body_text = " \n ".join(item["text"] for item in body_items)
    problems: list[str] = []

    header_paras = doc.sections[0].header.paragraphs
    if header_text:
        actual = header_paras[0].text if header_paras else ""
        if actual.strip() != header_text.strip():
            problems.append("promoted header text not found in doc.sections[0].header")
        if header_text in body_text:
            problems.append("promoted header text leaked into body")

    footer_paras = doc.sections[0].footer.paragraphs
    if footer_text:
        actual = footer_paras[0].text if footer_paras else ""
        if actual.strip() != footer_text.strip():
            problems.append("promoted footer text not found in doc.sections[0].footer")
        if footer_text in body_text:
            problems.append("promoted footer text leaked into body")

    return QACheck("header_footer_in_real_section", not problems, "; ".join(problems))


def _check_page_number_is_real_field(
    ast: DocumentAST, doc: Document, xml: str, mode: str
) -> QACheck:
    """
    Check that a verified page-number sequence becomes a real PAGE field.

    Never raw OCR'd digits standing in for one (acceptance criterion #6).
    Only enforced for mode="semantic" -- the visual renderer has no
    field-code support inside VML text boxes and is skipped, not failed.
    """
    if mode != "semantic":
        return QACheck(
            "page_number_uses_real_field", True, "skipped (visual mode has no fields)"
        )
    _header_plan, footer_plan = detect_header_footer_regions(ast)
    expects_page_field = any(r.has_page_field for r in footer_plan.regions) or (
        footer_plan.different_first_page
        and any(r.has_page_field for r in footer_plan.first_page_regions)
    )
    if not expects_page_field:
        return QACheck(
            "page_number_uses_real_field", True, "skipped (no page-number sequence)"
        )
    # Headers/footers live in their own XML parts (word/header1.xml etc.),
    # not inside doc.element.xml (the body) -- must be checked separately.
    header_xml = doc.sections[0].header._element.xml
    footer_xml = doc.sections[0].footer._element.xml
    found = bool(
        _INSTR_PAGE_RE.search(xml)
        or _INSTR_PAGE_RE.search(header_xml)
        or _INSTR_PAGE_RE.search(footer_xml)
    )
    detail = "" if found else "expected a real PAGE field, none found in output XML"
    return QACheck("page_number_uses_real_field", found, detail)
