"""Word Renderer — direct DOCX from Document AST with proper objects."""

from __future__ import annotations

import logging
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

from ..ast import DocumentAST, PageAST
from ..inline_markdown import parse_inline_segments
from ..latex_omml import LatexConversionError, latex_to_omml
from ..layout import LayoutType
from ..structure.header_footer import (
    HeaderFooterPlan,
    PromotedRegion,
    detect_header_footer_regions,
)
from ..structure.sections import SectionSpan, detect_sections
from .font_embed import embed_bundled_fonts
from .ooxml_helpers import add_field_run, add_hyperlink_run

logger = logging.getLogger(__name__)

FONT_LATIN = "Calibri"
FONT_CS = "IRNazanin"
FONT_MATH = "Cambria Math"

# Page size/margins fall back to these if the source page dims are missing
# or implausible (e.g. a non-PDF image source with no natural "page").
DEFAULT_PAGE_WIDTH_IN = 8.27  # A4
DEFAULT_PAGE_HEIGHT_IN = 11.69
MARGIN_RATIO = 0.08
MIN_MARGIN_IN = 0.4
MAX_MARGIN_IN = 1.0

# How close a node's bbox center must be to the page center, and how narrow
# it must be, to be treated as visually centered rather than a right-aligned
# block that merely doesn't reach the page edge.
CENTER_TOLERANCE_RATIO = 0.08
CENTER_MAX_WIDTH_RATIO = 0.7

MIN_IMAGE_WIDTH_IN = 1.5

# Unicode ranges of RTL scripts (Hebrew, Arabic incl. Persian core letters,
# Syriac, Arabic Supplement, Thaana, Arabic Extended-A, and the Hebrew/Arabic
# presentation-forms blocks). Identical to ``pdf.py``'s ``_RTL_RANGES``
# (kept in sync by hand, not imported) -- the sibling renderer uses these
# ranges for a *different* per-element heuristic (first-strong-character
# direction, HTML5 ``dir="auto"`` semantics), whereas this module needs a
# percentage-based ratio per the confirmed product rule below, so the two
# renderers share the underlying character classification but not the
# decision function built on top of it. Duplicated rather than extracted
# into a shared module because doing so would require editing pdf.py too,
# which is out of scope for this change.
_RTL_RANGES = (
    (0x0590, 0x05FF),  # Hebrew
    (0x0600, 0x06FF),  # Arabic (includes Persian core letters)
    (0x0700, 0x074F),  # Syriac
    (0x0750, 0x077F),  # Arabic Supplement
    (0x0780, 0x07BF),  # Thaana
    (0x08A0, 0x08FF),  # Arabic Extended-A
    (0xFB1D, 0xFDFF),  # Hebrew / Arabic presentation forms A
    (0xFE70, 0xFEFF),  # Arabic presentation forms B
)

# Confirmed product rule (user, verbatim): "if 20% or more of a paragraph's
# alphabetic characters are Persian/Arabic (RTL-range), treat the paragraph
# as RTL. Otherwise, treat it as LTR." The threshold is inclusive (>=),
# matching "20% or more" literally.
RTL_RATIO_THRESHOLD = 0.2


def _is_rtl_char(ch: str) -> bool:
    code = ord(ch)
    return any(start <= code <= end for start, end in _RTL_RANGES)


def _rtl_ratio(text: str) -> float:
    """
    Fraction of `text`'s *alphabetic* characters that are RTL-range.

    Digits, punctuation, whitespace, and symbols are excluded from both
    the numerator and denominator -- only alphabetic characters count, per
    the confirmed product rule. A paragraph with zero alphabetic
    characters (empty text, or text made up only of digits/punctuation/
    symbols -- e.g. a bare page number "7" or a rule "---") has no
    evidence of RTL content either way, so it defaults to LTR (ratio
    0.0) rather than RTL: that is both the conservative choice for the
    exact regression this fix targets (an all-English document must
    never end up forced RTL by non-alphabetic filler paragraphs) and
    consistent with English being python/Word's own "no information"
    default language.
    """
    alpha_chars = [ch for ch in text if ch.isalpha()]
    if not alpha_chars:
        return 0.0
    rtl_count = sum(1 for ch in alpha_chars if _is_rtl_char(ch))
    return rtl_count / len(alpha_chars)


def _is_rtl_paragraph(text: str) -> bool:
    """Return True when `text` meets the 20%-or-more RTL-alphabetic-char threshold."""
    return _rtl_ratio(text) >= RTL_RATIO_THRESHOLD


def render_docx(ast: DocumentAST, pdf_data: bytes | None = None) -> BytesIO:
    """
    Render full DocumentAST to a .docx BytesIO buffer.

    ``pdf_data`` (the original PDF bytes, when the source was a PDF) is
    used to detect the document's actual fonts instead of falling back to
    fixed Calibri/IRNazanin -- see ``detect_pdf_fonts`` in the legacy
    pipeline renderer, reused here rather than reimplemented.
    """
    font_cs, font_latin = _resolve_fonts(pdf_data)
    default_width_in, default_height_in = _resolve_page_size(ast)

    doc = Document()
    _modernize_compat_mode(doc)
    _set_theme_font_lang_bidi(doc)
    _setup_styles(doc, font_cs, font_latin)

    # Word is section-based: a landscape page (or any other page-size
    # change) partway through the document needs a real new section, not
    # just a page break at the first page's dimensions -- see
    # structure/sections.py.
    spans = detect_sections(ast)
    span_for_page = {
        n: span
        for span in spans
        for n in range(span.start_page_number, span.end_page_number + 1)
    }
    current_span = spans[0]
    content_width_in = _apply_section_size(
        doc.sections[0], current_span, default_width_in, default_height_in
    )

    if ast.title:
        doc.core_properties.title = ast.title

    header_plan, footer_plan = detect_header_footer_regions(ast)
    _apply_header_footer_plan(doc.sections[0], header_plan, is_header=True)
    _apply_header_footer_plan(doc.sections[0], footer_plan, is_header=False)
    header_promoted = header_plan.promoted
    footer_promoted = footer_plan.promoted
    page_number_promoted = any(r.has_page_field for r in footer_plan.regions) or (
        footer_plan.different_first_page
        and any(r.has_page_field for r in footer_plan.first_page_regions)
    )

    prev_page = 0
    for page in ast.pages:
        page_span = span_for_page.get(page.page_number, current_span)
        if prev_page > 0 and page.page_number > prev_page:
            if page_span is not current_span:
                new_section = doc.add_section(WD_SECTION.NEW_PAGE)
                content_width_in = _apply_section_size(
                    new_section, page_span, default_width_in, default_height_in
                )
                current_span = page_span
            else:
                doc.add_page_break()
        _render_page(
            doc,
            page,
            content_width_in,
            header_promoted=header_promoted,
            footer_promoted=footer_promoted,
            page_number_promoted=page_number_promoted,
        )
        prev_page = page.page_number

    buf = BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    # Embed the real font file(s) so the document doesn't depend on the
    # viewing device having font_cs installed -- a missing complex-script
    # font forces Word to substitute something else for it, and a
    # substitute without proper Arabic/Persian OpenType shaping tables can
    # misrender far beyond just the wrong glyph shapes. See font_embed.py.
    docx_bytes = embed_bundled_fonts(docx_bytes, {font_cs, font_latin, FONT_MATH})

    return BytesIO(docx_bytes)


def _modernize_compat_mode(doc: Document) -> None:
    """
    Bump compatibilityMode from python-docx's stale default of 14 to 15.

    python-docx's bundled blank template dates to Word 2010-era XML (compat
    mode 14); a document a human actually authored in a current Word never
    carries that legacy value. Word gates some layout-algorithm behavior on
    this setting, so leaving it at 14 is a real, avoidable difference from
    what real Word itself would produce -- unlike the schema-validated
    fixes elsewhere in this module, this isn't a confirmed defect, just one
    less place our output looks unlike a native Word file.
    """
    compat = doc.settings.element.find(qn("w:compat"))
    if compat is None:
        return
    for setting in compat.findall(qn("w:compatSetting")):
        if setting.get(qn("w:name")) == "compatibilityMode":
            setting.set(qn("w:val"), "15")


def _set_theme_font_lang_bidi(doc: Document) -> None:
    """
    Declare Persian as the document's own bidi/complex-script language.

    <w:themeFontLang> is a *document*-level language declaration in
    settings.xml, separate from the per-run/per-paragraph <w:lang> this
    module already sets everywhere -- python-docx's default template
    leaves it as plain en-US with no ``bidi`` attribute at all, meaning
    the document as a whole never declares itself as having any
    right-to-left content. Word's own internal decisions about how to
    treat RTL formatting can depend on this document-wide declaration,
    not just the per-paragraph one.
    """
    theme_font_lang = doc.settings.element.find(qn("w:themeFontLang"))
    if theme_font_lang is None:
        return
    theme_font_lang.set(qn("w:bidi"), "fa-IR")


def _set_doc_defaults_bidi_lang(doc: Document) -> None:
    """
    Fix the stray ar-SA default in styles.xml's docDefaults/rPrDefault.

    python-docx's template default is ``<w:lang ... w:bidi="ar-SA"/>`` --
    Arabic, not Persian. Every actual run/style/paragraph-mark lang is
    already overridden elsewhere in this module, so this default is never
    hit by real content, but it's a stray leftover pointing at the wrong
    language for a document-wide default, so it's corrected too.
    """
    doc_defaults = doc.styles.element.find(qn("w:docDefaults"))
    if doc_defaults is None:
        return
    rpr_default = doc_defaults.find(qn("w:rPrDefault") + "/" + qn("w:rPr"))
    if rpr_default is None:
        return
    lang = rpr_default.find(qn("w:lang"))
    if lang is not None:
        lang.set(qn("w:bidi"), "fa-IR")


def _resolve_fonts(pdf_data: bytes | None) -> tuple[str, str]:
    """Detect the source PDF's actual fonts; fall back to fixed defaults."""
    if not pdf_data:
        return FONT_CS, FONT_LATIN
    try:
        from ...pipeline.docx_renderer import detect_pdf_fonts

        detected = detect_pdf_fonts(pdf_data)
    except Exception:
        logger.debug("Font detection failed, using defaults", exc_info=True)
        return FONT_CS, FONT_LATIN
    return detected.get("cs", FONT_CS), detected.get("latin", FONT_LATIN)


def _resolve_page_size(ast: DocumentAST) -> tuple[float, float]:
    """
    Physical page size in inches, from the first page's rendered pixel.

    dimensions at its render DPI. Falls back to A4 if unavailable.
    """
    for page in ast.pages:
        if page.page_width > 0 and page.page_height > 0 and page.page_dpi > 0:
            return (
                page.page_width / page.page_dpi,
                page.page_height / page.page_dpi,
            )
    return DEFAULT_PAGE_WIDTH_IN, DEFAULT_PAGE_HEIGHT_IN


def _resolve_margin(page_width_in: float) -> float:
    return max(MIN_MARGIN_IN, min(MAX_MARGIN_IN, page_width_in * MARGIN_RATIO))


def _apply_header_footer_plan(
    section: object, plan: HeaderFooterPlan, *, is_header: bool
) -> None:
    """
    Write a detected header/footer plan onto a real Word section.

    One paragraph per promoted region (multiple regions == multiple
    independently-recurring elements, e.g. a logo line plus a separate
    banner line), and -- only when a genuine Different First Page variant
    was detected -- a separate set of paragraphs on section.first_page_*.
    """
    if not plan.promoted and not plan.different_first_page:
        return
    part = section.header if is_header else section.footer
    if plan.promoted:
        _write_promoted_regions(part, plan.regions)
    if plan.different_first_page:
        section.different_first_page_header_footer = True
        first_page_part = (
            section.first_page_header if is_header else section.first_page_footer
        )
        _write_promoted_regions(first_page_part, plan.first_page_regions)


def _write_promoted_regions(part: object, regions: tuple[PromotedRegion, ...]) -> None:
    """
    Write each region as its own paragraph with real PAGE/NUMPAGES fields.

    Reuses the part's existing first paragraph instead of always
    appending, so a call on an untouched header/footer part doesn't leave
    a stray leading empty paragraph.
    """
    part.is_linked_to_previous = False
    for i, region in enumerate(regions):
        reuse_first = i == 0 and part.paragraphs
        paragraph = part.paragraphs[0] if reuse_first else part.add_paragraph()
        paragraph.text = ""
        for segment in region.segments:
            if segment.kind == "text":
                if segment.text:
                    paragraph.add_run(segment.text)
            else:
                field_code = "PAGE" if segment.kind == "page" else "NUMPAGES"
                add_field_run(paragraph, field_code, segment.text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _apply_direction(paragraph, paragraph.text)


def _collect_header_footer(ast: DocumentAST) -> tuple[str, str]:
    """Pick the most common header/footer text across pages (a repeating banner)."""
    from collections import Counter

    headers = [
        node.text.strip()
        for page in ast.pages
        for node in page.nodes
        if node.type == LayoutType.header and node.text.strip()
    ]
    footers = [
        node.text.strip()
        for page in ast.pages
        for node in page.nodes
        if node.type == LayoutType.footer and node.text.strip()
    ]
    header = Counter(headers).most_common(1)[0][0] if headers else ""
    footer = Counter(footers).most_common(1)[0][0] if footers else ""
    return header, footer


def _set_section_header(doc: Document, text: str) -> None:
    section = doc.sections[0]
    section.header.is_linked_to_previous = False
    p = (
        section.header.paragraphs[0]
        if section.header.paragraphs
        else section.header.add_paragraph()
    )
    p.text = text
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_direction(p, text)


def _set_section_footer(doc: Document, text: str) -> None:
    section = doc.sections[0]
    section.footer.is_linked_to_previous = False
    p = (
        section.footer.paragraphs[0]
        if section.footer.paragraphs
        else section.footer.add_paragraph()
    )
    p.text = text
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_direction(p, text)


def _setup_styles(doc: Document, font_cs: str, font_latin: str) -> None:
    """
    Set up Normal/Heading styles, including their RTL bidi/alignment default.

    Deliberately still RTL by default (unchanged by the per-paragraph
    RTL-ratio fix below) -- every paragraph this renderer actually builds
    now runs through `_apply_direction`, which explicitly calls
    `_ensure_rtl` *or* `_ensure_ltr` on it based on its own text, so this
    style-level default no longer determines any real body content's
    direction; it only remains as the inherited fallback for the small
    number of paragraphs that carry no natural-language text of their own
    and already set their own alignment explicitly regardless of bidi
    (e.g. the centered OMML-equation container in `_add_formula`, the
    centered image container in `_add_image`). Flipping this document-wide
    default to LTR instead would be a much larger-blast-radius change --
    it would also flip the doc-defaults proofing language/complex-script
    settings for every one of those non-text containers -- for no
    behavioral benefit, since none of them have real text whose direction
    matters and this application's dominant use case (OCR of Persian
    documents) makes RTL the more sensible baseline for anything that
    *does* fall through to it. See `_ensure_ltr` for the paragraph-level
    override that actually fixes the reported all-English-document
    regression.
    """
    _set_doc_defaults_bidi_lang(doc)

    style = doc.styles["Normal"]
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _use_logical_alignment(style.element.get_or_add_pPr())

    # get_or_add_rFonts()/get_or_add_sz() (not a blind .append(), like this
    # code used to do) insert at their correct CT_RPr schema position --
    # rFonts must precede sz, and both must precede <w:lang/> -- regardless
    # of what's already present on rpr.
    rpr = style.element.get_or_add_rPr()
    r_fonts = rpr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_latin)
    r_fonts.set(qn("w:hAnsi"), font_latin)
    r_fonts.set(qn("w:cs"), font_cs)
    r_fonts.set(qn("w:eastAsia"), font_latin)
    sz = rpr.get_or_add_sz()
    if sz.get(qn("w:val")) is None:
        sz.set(qn("w:val"), "22")

    # A style's own <w:pPr> is CT_PPrGeneral, not CT_PPr -- unlike a real
    # paragraph's pPr, it has no nested <w:rPr> (paragraph-mark run props)
    # slot at all (confirmed against the ISO/IEC 29500-4 schema). The
    # style's own top-level <w:rPr> (below) is the only valid place for a
    # style's default run-level lang.
    p_pr = style.element.get_or_add_pPr()
    _insert_bidi(p_pr)
    _set_proof_lang(rpr)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hpf = hs.paragraph_format
        hpf.space_before = Pt(12)
        hpf.space_after = Pt(6)
        hpf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _use_logical_alignment(hs.element.get_or_add_pPr())
        _set_bidi_and_fonts(hs, font_cs, font_latin)


def _set_bidi_and_fonts(style: object, font_cs: str, font_latin: str) -> None:
    # <w:bidi/> is a *paragraph* property (CT_PPr/EG_PPrBase) -- it has no
    # valid place at all inside a style's <w:rPr> (confirmed against the
    # real ISO/IEC 29500-4 WordprocessingML schema: appending it there, as
    # this code used to do, produces XML with no valid parent for the
    # element anywhere in that content model, unlike a merely-misplaced-
    # but-elsewhere-valid element). It belongs on the style's own pPr,
    # exactly like the Normal style handles it in _setup_styles.
    p_pr = style.element.get_or_add_pPr()
    _insert_bidi(p_pr)

    # get_or_add_rFonts() (not a blind .append()) inserts at its correct
    # CT_RPr schema position regardless of what python-docx's own
    # font.bold/font.color setters already added to hpr.
    hpr = style.element.get_or_add_rPr()
    hr_fonts = hpr.get_or_add_rFonts()
    _set_proof_lang(hpr)
    hr_fonts.set(qn("w:cs"), font_cs)
    hr_fonts.set(qn("w:ascii"), font_latin)
    hr_fonts.set(qn("w:hAnsi"), font_latin)


def _apply_section_size(
    section: object,
    span: SectionSpan,
    default_width_in: float,
    default_height_in: float,
) -> float:
    """
    Apply a section's physical page size and return its usable content width.

    Falls back to the document default when the span's own size is
    unknown. The returned content width (in inches) is what the renderer
    uses to size/align content against.
    """
    width_in = span.page_width_in if span.page_width_in > 0 else default_width_in
    height_in = span.page_height_in if span.page_height_in > 0 else default_height_in
    margin_in = _resolve_margin(width_in)
    section.page_width = Inches(width_in)
    section.page_height = Inches(height_in)
    section.top_margin = Inches(margin_in)
    section.bottom_margin = Inches(margin_in)
    section.left_margin = Inches(margin_in)
    section.right_margin = Inches(margin_in)
    _insert_section_bidi(section._sectPr)
    _apply_column_count(section, span.column_count)
    content_width_in = max(1.0, width_in - 2 * margin_in)
    return content_width_in / max(1, span.column_count)


# CT_SectPrBase child elements that must come *after* <w:bidi/> per the
# OOXML schema (EG_SectPrContents sequence), in schema order.
_SECTPR_TAGS_AFTER_BIDI = (
    "w:rtlGutter",
    "w:docGrid",
    "w:printerSettings",
    "w:sectPrChange",
)


def _insert_section_bidi(sect_pr: object) -> None:
    """
    Mark a section itself as right-to-left, at its correct schema position.

    Per-paragraph <w:bidi/> (see _insert_bidi) is not enough on its own --
    observed directly in Microsoft Word for Mac: without this section-level
    flag, paragraphs with correct <w:bidi/> and <w:jc w:val="right"/> can
    still render pinned to the left margin. Idempotent, like _insert_bidi.
    """
    if sect_pr.find(qn("w:bidi")) is not None:
        return
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    sect_pr.insert_element_before(bidi, *_SECTPR_TAGS_AFTER_BIDI)


def _apply_column_count(section: object, column_count: int) -> None:
    """
    Set real Word section columns (w:cols).

    python-docx has no direct API for this, so the existing <w:cols>
    element (always present in a fresh section's sectPr) is updated in
    place; appending a second one would produce invalid, Word-rejected
    XML.
    """
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        # Defensive only (real sections always already have one) -- <w:cols/>
        # sits early in CT_SectPrBase, well before <w:bidi/>, so a blind
        # append here would misorder it behind bidi/docGrid if this branch
        # were ever actually hit on a section that already has those.
        cols = OxmlElement("w:cols")
        sect_pr.insert_element_before(
            cols,
            "w:formProt",
            "w:vAlign",
            "w:noEndnote",
            "w:titlePg",
            "w:textDirection",
            "w:bidi",
            "w:rtlGutter",
            "w:docGrid",
            "w:printerSettings",
            "w:sectPrChange",
        )
    cols.set(qn("w:num"), str(max(1, column_count)))


def _render_page(
    doc: Document,
    page: PageAST,
    content_width_in: float,
    header_promoted: bool = True,
    footer_promoted: bool = True,
    page_number_promoted: bool = True,
) -> None:
    for node in page.nodes:
        _render_node(
            doc,
            node,
            page.page_width,
            content_width_in,
            header_promoted,
            footer_promoted,
            page_number_promoted,
        )


_CAPTION_TYPES = (
    LayoutType.table_caption,
    LayoutType.table_footnote,
    LayoutType.figure_caption,
)


def _render_textish_node(
    doc: Document,
    node: object,
    page_width_px: float,
    header_promoted: bool = True,
    footer_promoted: bool = True,
    page_number_promoted: bool = True,
) -> bool:
    """Render heading/paragraph/list-like nodes; return True when handled."""
    if node.type == LayoutType.title:
        _add_heading(doc, node.text, 1, node, page_width_px)
        return True
    if node.type == LayoutType.heading:
        _add_heading(doc, node.text, min(node.level + 1, 3), node, page_width_px)
        return True
    if node.type == LayoutType.paragraph:
        _add_paragraph(doc, node.text, node=node, page_width_px=page_width_px)
        return True
    if node.type == LayoutType.reference:
        _add_paragraph(doc, f"📎 {node.text}", node=node, page_width_px=page_width_px)
        return True
    if node.type == LayoutType.list:
        _add_list(doc, node)
        return True
    if node.type in _CAPTION_TYPES:
        if node.text.strip():
            _add_paragraph(
                doc, node.text, italic=True, node=node, page_width_px=page_width_px
            )
        return True
    if node.type == LayoutType.unknown:
        if node.text.strip():
            _add_paragraph(doc, node.text, node=node, page_width_px=page_width_px)
        return True
    return _render_chrome_node(
        doc, node, page_width_px, header_promoted, footer_promoted, page_number_promoted
    )


def _render_chrome_node(
    doc: Document,
    node: object,
    page_width_px: float,
    header_promoted: bool,
    footer_promoted: bool,
    page_number_promoted: bool,
) -> bool:
    """
    Render header/footer/page-number nodes not promoted to real Word chrome.

    A header/footer/page-number sequence that didn't clear the cross-page
    repetition bar (e.g. a single-page document) or couldn't be verified
    as real is page-specific content, not chrome -- it still needs to
    render somewhere rather than being silently dropped, but is never
    promoted/faked as real Word chrome either way.
    """
    promoted_by_type = {
        LayoutType.header: header_promoted,
        LayoutType.footer: footer_promoted,
        LayoutType.page_number: page_number_promoted,
    }
    if node.type not in promoted_by_type:
        return False
    if not promoted_by_type[node.type] and node.text.strip():
        _add_paragraph(doc, node.text, node=node, page_width_px=page_width_px)
    return True


def _render_media_node(
    doc: Document, node: object, page_width_px: float, content_width_in: float
) -> bool:
    """Render table/formula/figure/chart/code; return True when handled."""
    if node.type == LayoutType.table:
        _add_table(doc, node)
        return True
    if node.type == LayoutType.formula:
        _add_formula(doc, node.latex)
        return True
    if node.type == LayoutType.figure:
        _add_image(doc, node, page_width_px, content_width_in)
        return True
    if node.type == LayoutType.chart:
        _add_chart(doc, node, page_width_px, content_width_in)
        return True
    if node.type == LayoutType.code:
        _add_code(doc, node)
        return True
    return False


def _render_node(
    doc: Document,
    node: object,
    page_width_px: float,
    content_width_in: float,
    header_promoted: bool = True,
    footer_promoted: bool = True,
    page_number_promoted: bool = True,
) -> None:
    if _render_textish_node(
        doc, node, page_width_px, header_promoted, footer_promoted, page_number_promoted
    ):
        return
    _render_media_node(doc, node, page_width_px, content_width_in)


def _resolve_alignment(node: object, page_width_px: float) -> object:
    """
    Right-align by default (RTL body text); promote to CENTER when centered.

    "Centered" means the node's bbox is genuinely centered on the page
    rather than merely falling short of the page edge (e.g. a centered
    title/box heading).
    """
    if not page_width_px or node.bbox == (0.0, 0.0, 0.0, 0.0):
        return WD_ALIGN_PARAGRAPH.RIGHT
    x1, _, x2, _ = node.bbox
    width = x2 - x1
    center_offset = abs((x1 + x2) / 2 - page_width_px / 2)
    if (
        width < page_width_px * CENTER_MAX_WIDTH_RATIO
        and center_offset < page_width_px * CENTER_TOLERANCE_RATIO
    ):
        return WD_ALIGN_PARAGRAPH.CENTER
    return WD_ALIGN_PARAGRAPH.RIGHT


def _add_heading(
    doc: Document,
    text: str,
    level: int,
    node: object = None,
    page_width_px: float = 0.0,
) -> None:
    level = max(1, min(level, 3))
    p = doc.add_heading(text, level=level)
    p.alignment = (
        _resolve_alignment(node, page_width_px) if node else WD_ALIGN_PARAGRAPH.RIGHT
    )
    _apply_direction(p, p.text)


def _add_inline_runs(paragraph: object, text: str) -> None:
    """
    Add one styled run/hyperlink/inline-math object per inline-Markdown span.

    A real run per bold/italic/code span, a real ``w:hyperlink`` for a
    link span, or a real inline ``m:oMath`` for a dollar-delimited math
    span -- instead of dumping raw
    ``**text**``/`` `code` ``/``[text](url)``/``$latex$`` markers into a
    single plain run. See ../inline_markdown.py.
    """
    for seg in parse_inline_segments(text):
        if not seg.text:
            continue
        if seg.url:
            add_hyperlink_run(paragraph, seg.text, seg.url)
            continue
        if seg.math:
            _add_inline_math(paragraph, seg.text)
            continue
        run = paragraph.add_run(seg.text)
        run.bold = seg.bold
        run.italic = seg.italic
        if seg.code:
            run.font.name = "Courier New"


def _add_inline_math(paragraph: object, latex: str) -> None:
    """
    Insert a real inline ``m:oMath`` equation object mid-paragraph.

    Instead of leaving raw ``$latex$`` sitting inside RTL text -- mixed
    Persian/Latin/math tokens with no directional isolation reliably
    produce bidi word-reordering artifacts (observed directly rendering a
    real OCR'd document), whereas a real OMML object is one atomic inline
    element Word positions correctly regardless of the surrounding
    paragraph's direction. Falls back to styled math-font text, matching
    the block-formula fallback in _add_formula, if the LaTeX can't be
    parsed.
    """
    try:
        body = latex_to_omml(latex)  # "<m:oMath>...</m:oMath>", no ns decl
        omath_xml = body.replace("<m:oMath>", f"<m:oMath {nsdecls('m')}>", 1)
        paragraph._p.append(parse_xml(omath_xml))
    except LatexConversionError:
        logger.warning(
            "Inline LaTeX->OMML conversion failed, using text fallback: %r", latex
        )
        run = paragraph.add_run(latex)
        run.italic = True
        run.font.name = FONT_MATH


def _add_paragraph(
    doc: Document,
    text: str,
    italic: bool = False,
    node: object = None,
    page_width_px: float = 0.0,
) -> None:
    p = doc.add_paragraph()
    if italic:
        run = p.add_run(text or "")
        run.italic = True
    else:
        _add_inline_runs(p, text or "")
    p.alignment = (
        _resolve_alignment(node, page_width_px) if node else WD_ALIGN_PARAGRAPH.RIGHT
    )
    _apply_direction(p, p.text)


def _add_list(doc: Document, node: object) -> None:
    for child in node.children:
        style = "List Number" if getattr(child, "ordered", False) else "List Bullet"
        p = doc.add_paragraph(style=style)
        _add_inline_runs(p, child.text or "")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _apply_direction(p, p.text)


def _add_table(doc: Document, node: object) -> None:
    if not node.rows:
        return
    rows = node.rows
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            if ci >= cols:
                break
            cell = table.cell(ri, ci)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            _add_inline_runs(paragraph, str(cell_text))
            # Ratio computed per cell, not per table/row: a genuinely mixed
            # table (e.g. an English label column next to a Persian value
            # column) needs each cell judged on its own content, not one
            # table-wide verdict that would force every cell into
            # whichever language happens to dominate the table overall.
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _apply_direction(paragraph, paragraph.text)

    # Reconstruct rowspan/colspan as real merged Word table cells instead of
    # leaving them as separate cells with duplicated/blank text -- see
    # ast.py's _parse_html_table for how these spans are recovered from the
    # source HTML.
    for r1, c1, r2, c2 in node.cell_merges:
        if r2 < len(rows) and c2 < cols:
            table.cell(r1, c1).merge(table.cell(r2, c2))

    if getattr(node, "repeat_header_row", False) and rows:
        _mark_row_as_repeating_header(table.rows[0])


def _mark_row_as_repeating_header(row: object) -> None:
    """
    Set w:tblHeader on a row so Word repeats it on every printed page.

    See structure/table_continuation.py, which produces multi-page-source
    tables long enough to need this.
    """
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))


def _add_formula(doc: Document, latex: str) -> None:
    """
    Add formula as a real, editable OMML (Office Math) Word equation.

    LaTeX is parsed and re-emitted as genuine ``m:oMath`` XML (see
    ../latex_omml.py) so Word treats it as an equation object, not text.
    If the LaTeX can't be parsed, fall back to styled math-font text
    instead of failing the whole document.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        omath_xml = latex_to_omml(latex)
        o_math_para = parse_xml(
            f"<m:oMathPara {nsdecls('m')}>{omath_xml}</m:oMathPara>"
        )
        p._p.append(o_math_para)
    except LatexConversionError:
        logger.warning(
            "Formula LaTeX->OMML conversion failed, using text fallback: %r", latex
        )
        run = p.add_run(latex)
        run.italic = True
        run.font.name = FONT_MATH


def _resolve_image_width_in(
    node: object, page_width_px: float, content_width_in: float
) -> float:
    """
    Size the image relative to how much of the page width it occupied.

    Uses the original bbox width instead of always inserting a fixed 5
    inches, so a small inline figure doesn't dominate the page and a
    full-width diagram isn't shrunk down to match everything else.
    """
    if not page_width_px or node.bbox == (0.0, 0.0, 0.0, 0.0):
        return min(5.0, content_width_in)
    x1, _, x2, _ = node.bbox
    ratio = max(0.0, min(1.0, (x2 - x1) / page_width_px))
    width_in = ratio * content_width_in
    return max(MIN_IMAGE_WIDTH_IN, min(width_in, content_width_in))


def _add_image(
    doc: Document,
    node: object,
    page_width_px: float = 0.0,
    content_width_in: float = 6.0,
) -> None:
    asset_path = node.asset_path
    if not asset_path or not Path(asset_path).exists():
        _add_paragraph(doc, f"[تصویر: {node.caption}]", italic=True)
        return

    try:
        img_bytes = Path(asset_path).read_bytes()
        width_in = _resolve_image_width_in(node, page_width_px, content_width_in)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(BytesIO(img_bytes), width=Inches(width_in))
        alt_text = " ".join(t for t in (node.caption, node.description) if t).strip()
        if alt_text:
            _set_image_alt_text(run, alt_text)
    except Exception:
        _add_paragraph(doc, f"[تصویر: {node.caption}]", italic=True)


def _set_image_alt_text(run: object, alt_text: str) -> None:
    """
    Set an inline picture's real Word "Alt Text" (accessibility description).

    ``node.caption``/``node.description`` are both VLM-generated -- the
    VLM is asked for a short caption vs. a longer description, but there
    is no reliable way to tell its "short caption" apart from just a
    terser description (e.g. "Young Scholars Club logo" *is* a caption
    by that prompt's own definition, yet is exactly the same kind of
    content a sighted reader doesn't need duplicated as a paragraph
    under the image). Any real caption actually printed in the source
    document is captured independently by the layout detector as its
    own text element (rendered normally, like any other paragraph) --
    so neither VLM field is ever shown as visible body text here, both
    go into <wp:docPr descr="..."/>, the same field Word's own "Edit Alt
    Text" panel reads/writes.
    """
    doc_pr = run._element.find(f".//{qn('wp:docPr')}")
    if doc_pr is not None:
        doc_pr.set("descr", alt_text)


def _add_chart(
    doc: Document,
    node: object,
    page_width_px: float = 0.0,
    content_width_in: float = 6.0,
) -> None:
    """Render chart — image (with alt text) plus optional data table."""
    _add_image(doc, node, page_width_px, content_width_in)
    if node.chart_data and node.chart_data.get("data"):
        rows = [["Label", "Value"]]
        rows.extend(
            [str(item.get("label", "")), str(item.get("value", ""))]
            for item in node.chart_data["data"]
            if isinstance(item, dict)
        )
        tbl = doc.add_table(rows=len(rows), cols=2)
        tbl.style = "Table Grid"
        for ri, row_data in enumerate(rows):
            for ci, val in enumerate(row_data):
                tbl.cell(ri, ci).text = val


def _add_code(doc: Document, node: object) -> None:
    p = doc.add_paragraph()
    run = p.add_run(node.text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    # Code is LTR content regardless of the surrounding RTL document -- left
    # unset, this paragraph would silently inherit the Normal style's
    # RIGHT alignment (see _setup_styles), right-aligning source code.
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_shading(p)


# CT_PPr child elements that must come *after* <w:bidi/> per the OOXML
# schema (EG_PPrBase sequence), in schema order. <w:jc/> (paragraph
# alignment/justification) is among them: appending <w:bidi/> at the end of
# an already-built pPr -- as this code used to do -- puts it after <w:jc/>,
# violating that order. Real Microsoft Word can be strict about CT_PPr
# element order and silently drop/ignore misordered properties on open
# (observed directly: a paragraph with correct WD_ALIGN_PARAGRAPH.RIGHT set
# in python-docx, and correct rendering in LibreOffice, showing up
# left-aligned in real Word) -- LibreOffice/python-docx are lenient readers
# and don't surface this, which is why it went unnoticed until tested in
# real Word.
_PPR_TAGS_AFTER_BIDI = (
    "w:adjustRightInd",
    "w:snapToGrid",
    "w:spacing",
    "w:ind",
    "w:contextualSpacing",
    "w:mirrorIndents",
    "w:suppressOverlap",
    "w:jc",
    "w:textDirection",
    "w:textAlignment",
    "w:textboxTightWrap",
    "w:outlineLvl",
    "w:divId",
    "w:cnfStyle",
    "w:rPr",
    "w:sectPr",
    "w:pPrChange",
)

# <w:shd/> sits earlier still in the same CT_PPrBase sequence (right after
# <w:pBdr/>) -- reuses _PPR_TAGS_AFTER_BIDI for everything from <w:bidi/>
# onward rather than repeating that tail a second time.
_PPR_TAGS_AFTER_SHD = (
    "w:tabs",
    "w:suppressAutoHyphens",
    "w:kinsoku",
    "w:wordWrap",
    "w:overflowPunct",
    "w:topLinePunct",
    "w:autoSpaceDE",
    "w:autoSpaceDN",
    "w:bidi",
    *_PPR_TAGS_AFTER_BIDI,
)


def _set_shading(p: object) -> None:
    p_pr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    p_pr.insert_element_before(shd, *_PPR_TAGS_AFTER_SHD)


def _apply_direction(p: object, text: str) -> None:
    """
    Pick RTL vs. LTR for paragraph `p` from `text`'s own RTL-alphabetic ratio.

    This is the single decision point every call site funnels through --
    replaces what used to be an unconditional `_ensure_rtl(p)` call at
    every paragraph/heading/list-item/table-cell/header-footer site,
    which forced RTL bidi + right alignment onto every paragraph
    regardless of its actual language content (the reported regression:
    an all-English source document rendered as an all-RTL Word file).
    """
    if _is_rtl_paragraph(text):
        _ensure_rtl(p)
    else:
        _ensure_ltr(p)


def _ensure_rtl(p: object) -> None:
    p_pr = p._element.get_or_add_pPr()
    _insert_bidi(p_pr, rtl=True)
    _use_logical_alignment(p_pr)
    _set_proof_lang(_get_or_add_ppr_rpr(p_pr))
    for run in p.runs:
        _set_proof_lang(run._element.get_or_add_rPr())


def _ensure_ltr(p: object) -> None:
    """
    Explicitly mark a paragraph as left-to-right, overriding the RTL style default.

    See `_setup_styles`/`_set_bidi_and_fonts` for that default. Mirrors
    `_ensure_rtl`'s structure and OOXML-schema-order discipline
    (same `_insert_bidi` insertion helper, same schema-order tag list,
    same paragraph-mark/`run` `rPr` handling), but differs in three ways
    that matter:

    - `_insert_bidi(p_pr, rtl=False)` writes an *explicit*
      `<w:bidi w:val="0"/>`, not merely an absent `<w:bidi/>`. Absent
      would silently fall through to the style's own `<w:bidi w:val="1"/>`
      default (see `_setup_styles`) rather than actually flipping this
      paragraph to LTR -- deliberately left at the style level rather than
      changed there (see the module-level note above `_setup_styles`), so
      every paragraph that isn't judged RTL must positively override it.
    - Alignment: `_use_logical_alignment` (the RTL branch) rewrites a
      physical `<w:jc w:val="right"/>` to the *logical* "start" value so it
      keeps meaning "right" once `<w:bidi/>` flips the paragraph's reading
      direction (see that function's docstring for why "right" itself
      isn't enough in real Word). An LTR paragraph has no such
      reinterpretation to guard against -- its reading-direction start is
      already the physical left edge -- so `_use_physical_left_alignment`
      instead rewrites that same "right" default (this renderer's
      RTL-body-text assumption, set before either `_ensure_*` function
      runs) to plain physical "left". Center/other `<w:jc/>` values (a
      genuinely bbox-centered heading, a centered header/footer line) are
      left untouched either way.
    - Language: `_set_ltr_lang` sets `<w:rtl w:val="0"/>` and a plain
      `en-US` `<w:lang/>` (no `w:bidi` attribute) rather than Persian, so
      spell-check and other language-dependent Word behavior treats this
      paragraph as ordinary Latin-script text.
    """
    p_pr = p._element.get_or_add_pPr()
    _insert_bidi(p_pr, rtl=False)
    _use_physical_left_alignment(p_pr)
    _set_ltr_lang(_get_or_add_ppr_rpr(p_pr))
    for run in p.runs:
        _set_ltr_lang(run._element.get_or_add_rPr())


def _use_logical_alignment(p_pr: object) -> None:
    """
    Convert a physical <w:jc w:val="right"/> to the logical "start" value.

    Per spec, "right"/"left" are physical (unaffected by <w:bidi/>) while
    "start"/"end" are logical (flip with paragraph direction) -- so both
    should produce the same result for an RTL paragraph. They don't, in
    real Microsoft Word: confirmed directly by the user testing generated
    files side by side (Word for Mac + mobile) -- "right" rendered
    left-aligned despite correct <w:bidi/>, "start" rendered correctly
    right-aligned. <w:jc>'s other values (center, both, ...) are left
    untouched; only paragraphs actually using physical "right" (this
    renderer's stand-in for "the reading-direction-start alignment") are
    affected.
    """
    jc = p_pr.find(qn("w:jc"))
    if jc is not None and jc.get(qn("w:val")) == "right":
        jc.set(qn("w:val"), "start")


def _use_physical_left_alignment(p_pr: object) -> None:
    """
    Convert a physical <w:jc w:val="right"/> to physical "left", for LTR paragraphs.

    Mirrors `_use_logical_alignment`'s role for the RTL branch, but the
    correct target value is different: an RTL paragraph needs the
    *logical* "start" value (see that function's docstring for why plain
    "right" isn't enough in real Word), whereas an LTR paragraph's
    reading-direction start is already the physical left edge -- plain
    "left" is correct and unambiguous, no bidi-dependent logical value
    needed. <w:jc>'s other values (center, both, ...) -- e.g. a genuinely
    bbox-centered heading, or a centered header/footer paragraph -- are
    left untouched, exactly like `_use_logical_alignment` leaves them.
    """
    jc = p_pr.find(qn("w:jc"))
    if jc is not None and jc.get(qn("w:val")) == "right":
        jc.set(qn("w:val"), "left")


def _insert_bidi(p_pr: object, *, rtl: bool = True) -> None:
    """
    Set <w:bidi/> on a pPr at its correct schema position (idempotent).

    Explicit w:val="1"/"0" rather than a bare <w:bidi/> (for the RTL case)
    or simply leaving it absent (for the LTR case) -- CT_OnOff's implied
    default is "true" per spec, but not every Word build reliably applies
    that implicit default; a known-working reference implementation
    (hawzeh-tts's docx_lang.py/markdown_docx.py) always sets it
    explicitly. An explicit "0" additionally matters here because the
    Normal/Heading *styles* this paragraph inherits from carry their own
    `<w:bidi w:val="1"/>` (see `_setup_styles`/`_set_bidi_and_fonts`) -- an
    absent paragraph-level <w:bidi/> would silently fall through to that
    RTL style default instead of actually overriding it to LTR.

    Idempotent for repeated calls with the *same* `rtl` value (matching
    the original behavior); a later call with a *different* `rtl` value
    updates the existing element's value in place rather than being a
    no-op, so re-deciding a paragraph's direction after an earlier call
    still lands correctly without reordering anything already inserted.
    """
    val = "1" if rtl else "0"
    existing = p_pr.find(qn("w:bidi"))
    if existing is not None:
        existing.set(qn("w:val"), val)
        return
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), val)
    p_pr.insert_element_before(bidi, *_PPR_TAGS_AFTER_BIDI)


def _get_or_add_ppr_rpr(p_pr: object) -> object:
    """
    Get/create the paragraph-mark's own <w:rPr> nested inside a <w:pPr>.

    Distinct from a run's <w:rPr>: this is CT_PPr's own optional rPr
    child (the paragraph mark's formatting), which per the OOXML schema
    must be the last real content child, right before sectPr/pPrChange.
    <w:lang> (like other run-level formatting, e.g. <w:rFonts>) is only
    valid here nested inside this element -- never as a direct child of
    <w:pPr> itself, which is what a previous version of this code did.
    """
    existing = p_pr.find(qn("w:rPr"))
    if existing is not None:
        return existing
    rpr = OxmlElement("w:rPr")
    p_pr.insert_element_before(rpr, "w:sectPr", "w:pPrChange")
    return rpr


def _set_proof_lang(pr: object) -> None:
    """
    Set proofing language and right-to-left script on an rPr element.

    ``pr`` must be a real <w:rPr> -- a run's, a style's, or a paragraph
    mark's own nested one (see _get_or_add_ppr_rpr) -- never a <w:pPr>
    directly: <w:lang> is not a valid direct child of CT_PPr, and real
    Microsoft Word can silently drop/repair a paragraph's formatting on
    open when it is (LibreOffice/python-docx tolerate it and don't
    surface the problem, which is how this went unnoticed before).

    Also sets <w:rtl/>, a distinct *run-level* property from the
    *paragraph-level* <w:bidi/> set by _insert_bidi. This is exactly what
    real Word itself writes when a human selects text and clicks
    "Right-to-left text direction" -- without it, paragraph-level bidi
    alone was not always enough for real Word to lay the paragraph out
    against the right margin, even with correct <w:jc w:val="right"/>.
    """
    rtl = pr.get_or_add_rtl()
    rtl.set(qn("w:val"), "1")
    for old in pr.findall(qn("w:lang")):
        pr.remove(old)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "fa-IR")
    lang.set(qn("w:bidi"), "fa-IR")
    pr.append(lang)


def _set_ltr_lang(pr: object) -> None:
    """
    Set proofing language and left-to-right script on an rPr element.

    The LTR mirror of `_set_proof_lang` -- see that function's docstring
    for why `pr` must be a real <w:rPr> and never a <w:pPr> directly.

    Sets <w:rtl w:val="0"/> (explicit, for the same reason `_insert_bidi`
    sets an explicit <w:bidi w:val="0"/> rather than leaving it absent:
    the surrounding style default is RTL) and a plain `en-US` <w:lang/>
    with no `w:bidi` attribute. The `w:bidi` attribute on `<w:lang/>` only
    names the document's complex-script *proofing* language for whatever
    complex-script runs might appear -- it doesn't affect this run's own
    reading direction (that's `w:rtl`/`w:bidi` above) -- so it's simply
    left unset here rather than pointed at some LTR-appropriate value,
    matching how real Word itself leaves it off a plain Latin-script run.
    """
    rtl = pr.get_or_add_rtl()
    rtl.set(qn("w:val"), "0")
    for old in pr.findall(qn("w:lang")):
        pr.remove(old)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "en-US")
    pr.append(lang)
