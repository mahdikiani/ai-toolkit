"""Deterministic DOCX renderer — builds a Word document from pipeline output."""

from __future__ import annotations

import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage

logger = logging.getLogger(__name__)

FONT_PERSIAN = "B Nazanin"
FONT_LATIN = "Calibri"


def _tally_pdf_fonts(pdf_data: bytes) -> tuple[dict[str, int], dict[str, int]]:
    import fitz

    latin_fonts: dict[str, int] = {}
    cs_fonts: dict[str, int] = {}
    try:
        doc = fitz.open(stream=pdf_data, filetype="pdf")
        for page in doc:
            for block in page.get_text("dict").get("blocks", []):
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        name = span.get("font", "")
                        if not name:
                            continue
                        is_cs = bool(span.get("bidi") or 0) or any(
                            ord(c) > 0x0600 for c in span.get("text", "")[:10]
                        )
                        bucket = cs_fonts if is_cs else latin_fonts
                        bucket[name] = bucket.get(name, 0) + 1
        doc.close()
    except Exception:
        logger.debug("Could not read fonts from PDF", exc_info=True)
    return latin_fonts, cs_fonts


def _pick_cleaned_font(counts: dict[str, int]) -> str | None:
    if not counts:
        return None
    return _clean_font_name(max(counts, key=counts.get))


def detect_pdf_fonts(pdf_data: bytes | None = None) -> dict[str, str]:
    """
    Extract most common Persian and Latin font names from PDF spans.

    Only returns font names that look like real system fonts
    (strips foundry prefixes like XB, B, I, R and filters numeric/symbolic names).
    """
    if not pdf_data:
        return {}
    latin_fonts, cs_fonts = _tally_pdf_fonts(pdf_data)
    result: dict[str, str] = {}
    clean_cs = _pick_cleaned_font(cs_fonts)
    if clean_cs:
        result["cs"] = clean_cs
    clean_latin = _pick_cleaned_font(latin_fonts)
    if clean_latin:
        result["latin"] = clean_latin
    return result


def _clean_font_name(name: str) -> str | None:
    """
    Clean a PDF-internal font name to find the real system font name.

    Strips foundry prefixes (XB, B, I, R, etc.), filters names with digits
    or special characters that can't be real system fonts.
    """
    import re

    name = name.strip()
    if not name:
        return None

    if len(name) < 3:
        return None

    # Reject purely numeric/symbolic names like CMMI12, CMR10
    if re.match(r"^[A-Z]+[\d]+$", name):
        return None
    if name.startswith("CM") and any(c.isdigit() for c in name):
        return None

    # Strip common foundry prefixes
    foundry_prefixes = r"^(XB|IB|IR|B|I|R|L|M)[\s-]*"
    cleaned = re.sub(foundry_prefixes, "", name, flags=re.IGNORECASE).strip()

    # Common font name mappings
    name_lower = cleaned.lower()
    mappings = {
        "niloofar": "B Nazanin",
        "b nazanin": "B Nazanin",
        "bnazanin": "B Nazanin",
        "nazanin": "B Nazanin",
        "lotus": "B Lotus",
        "mitra": "B Mitra",
        "yekan": "B Yekan",
        "zar": "B Zar",
        "traffic": "B Traffic",
        "roya": "B Roya",
        "nassim": "B Nassim",
        "times": "Times New Roman",
        "arial": "Arial",
        "tahoma": "Tahoma",
        "calibri": "Calibri",
        "cambria": "Cambria",
        "timesnewroman": "Times New Roman",
    }

    for key, value in mappings.items():
        if key in name_lower:
            return value

    return None


def build_docx(
    markdown: str = "",
    *,
    page_images: list[PILImage.Image] | None = None,
    elements: list[Any] | None = None,
    page_headers: dict[int, str] | None = None,
    page_footers: dict[int, str] | None = None,
    crops: dict[str, bytes] | None = None,
    crops_dir: str | Path | None = None,
    assets_dir: str | Path | None = None,
    pdf_data: bytes | None = None,
) -> BytesIO:
    """
    Convert pipeline output to a .docx BytesIO.

    Supports: headings, paragraphs, lists, images, tables, formulas,
    headers/footers, RTL text, and font detection.
    """
    doc = Document()

    # ── Fonts ────────────────────────────────────────────────────────────────
    detected = detect_pdf_fonts(pdf_data)
    cs_font = detected.get("cs") or FONT_PERSIAN
    latin_font = detected.get("latin") or FONT_LATIN
    logger.info("DOCX fonts — Persian: %s, English: %s", cs_font, latin_font)

    # ── Page setup ───────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    _setup_rtl_styles(doc, cs_font=cs_font, latin_font=latin_font)

    # ── Headers / Footers ────────────────────────────────────────────────────
    page_count = max([
        len(page_images or []),
        max(page_headers.keys()) if page_headers else 0,
        max(page_footers.keys()) if page_footers else 0,
        0,
    ])
    if page_count > 1:
        _add_page_headers_footers(
            doc, page_headers or {}, page_footers or {}, cs_font, latin_font, page_count
        )

    # ── Markdown body ────────────────────────────────────────────────────────
    if markdown.strip():
        blocks = _parse_markdown_blocks(markdown)
        for block in blocks:
            _render_block(
                doc, block, page_images, elements, crops, crops_dir, assets_dir
            )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_page_headers_footers(
    doc: Document,
    headers: dict[int, str],
    footers: dict[int, str],
    cs_font: str,
    latin_font: str,
    total_pages: int,
) -> None:
    """Add per-page headers and footers using section breaks."""
    section = doc.sections[0]
    _set_section_header(section, headers.get(1, ""), cs_font, latin_font, is_first=True)
    _set_section_footer(section, footers.get(1, ""), cs_font, latin_font, is_first=True)

    for page_num in range(2, total_pages + 1):
        doc.add_page_break()
        # Add section break for each page's header/footer
        new_section = doc.add_section()
        _set_section_header(new_section, headers.get(page_num, ""), cs_font, latin_font)
        _set_section_footer(new_section, footers.get(page_num, ""), cs_font, latin_font)


def _set_section_header(
    section: object, text: str, cs_font: str, latin_font: str, is_first: bool = False
) -> None:
    if not text:
        return
    header = section.header
    p = header.add_paragraph() if not header.paragraphs else header.paragraphs[0]
    p.text = ""
    run = p.add_run(text)
    run.font.name = cs_font
    _set_run_fonts(run, cs_font, latin_font)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _set_section_footer(
    section: object, text: str, cs_font: str, latin_font: str, is_first: bool = False
) -> None:
    if not text:
        return
    footer = section.footer
    p = footer.add_paragraph() if not footer.paragraphs else footer.paragraphs[0]
    p.text = ""
    run = p.add_run(text)
    run.font.name = cs_font
    _set_run_fonts(run, cs_font, latin_font)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── Styling ──────────────────────────────────────────────────────────────────


def _setup_rtl_styles(
    doc: Document,
    *,
    cs_font: str = FONT_PERSIAN,
    latin_font: str = FONT_LATIN,
) -> None:
    style = doc.styles["Normal"]
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _xml(tag: str, **attrs: str) -> OxmlElement:
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(qn(k), v)
        return el

    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))  # ruff: ignore[non-lowercase-variable-in-function]
    if rFonts is None:
        rFonts = _xml("w:rFonts")  # ruff: ignore[non-lowercase-variable-in-function]
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), latin_font)
    rFonts.set(qn("w:hAnsi"), latin_font)
    rFonts.set(qn("w:cs"), cs_font)

    pPr = style.element.get_or_add_pPr()  # ruff: ignore[non-lowercase-variable-in-function]
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        pPr.append(_xml("w:bidi"))

    sz = rpr.find(qn("w:sz"))
    if sz is None:
        rpr.append(_xml("w:sz", **{"w:val": "22"}))

    for level in range(1, 5):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = cs_font
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0, 0, 0)
        pf_h = hs.paragraph_format
        pf_h.space_before = Pt(12)
        pf_h.space_after = Pt(6)
        pf_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hpr = hs.element.get_or_add_rPr()
        if hpr.find(qn("w:bidi")) is None:
            hpr.append(_xml("w:bidi"))
        hrFonts = hpr.find(qn("w:rFonts"))  # ruff: ignore[non-lowercase-variable-in-function]
        if hrFonts is None:
            hrFonts = _xml("w:rFonts")  # ruff: ignore[non-lowercase-variable-in-function]
            hpr.append(hrFonts)
        hrFonts.set(qn("w:cs"), cs_font)
        hrFonts.set(qn("w:ascii"), latin_font)
        hrFonts.set(qn("w:hAnsi"), latin_font)


def _set_run_fonts(run: object, cs_font: str, latin_font: str) -> None:
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))  # ruff: ignore[non-lowercase-variable-in-function]
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")  # ruff: ignore[non-lowercase-variable-in-function]
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), latin_font)
    rFonts.set(qn("w:hAnsi"), latin_font)
    rFonts.set(qn("w:cs"), cs_font)
    rFonts.set(qn("w:eastAsia"), latin_font)


def _set_run_fonts_from_doc(doc: Document, run: object) -> None:
    """Apply Normal style fonts to a run element."""
    normal_style = doc.styles["Normal"]
    rpr = normal_style.element.find(qn("w:rPr"))
    if rpr is None:
        return
    rFonts = rpr.find(qn("w:rFonts"))  # ruff: ignore[non-lowercase-variable-in-function]
    if rFonts is None:
        return
    run_rpr = run._element.get_or_add_rPr()
    run_rFonts = run_rpr.find(qn("w:rFonts"))  # ruff: ignore[non-lowercase-variable-in-function]
    if run_rFonts is None:
        run_rFonts = OxmlElement("w:rFonts")  # ruff: ignore[non-lowercase-variable-in-function]
        run_rpr.append(run_rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        val = rFonts.get(qn(attr))
        if val:
            run_rFonts.set(qn(attr), val)


# ── Block rendering ──────────────────────────────────────────────────────────


def _render_block(
    doc: Document,
    block: dict,
    page_images: list[PILImage.Image] | None,
    elements: list[Any] | None,
    crops: dict[str, bytes] | None,
    crops_dir: str | Path | None,
    assets_dir: str | Path | None,
) -> None:
    kind = block.get("kind")
    if kind == "heading":
        _add_heading(doc, block["text"], block["level"])
    elif kind == "paragraph":
        _add_paragraph(doc, block["text"])
    elif kind == "list_item":
        _add_paragraph(doc, f"• {block['text']}")
    elif kind == "image":
        _add_image_block(
            doc, block, page_images, elements, crops, crops_dir, assets_dir
        )
    elif kind == "table":
        _add_table(doc, block)
    elif kind == "formula":
        _add_formula(doc, block["latex"], block.get("display", False))
    elif kind == "separator":
        doc.add_page_break()
    elif kind == "blockquote":
        _add_paragraph(doc, block["text"], italic=True)


def _add_heading(doc: Document, text: str, level: int) -> None:
    level = max(1, min(level, 4))
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _ensure_rtl(p)


def _add_paragraph(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _ensure_rtl(p)
    _set_run_fonts_from_doc(doc, run)


def _add_image_block(
    doc: Document,
    block: dict,
    page_images: list[PILImage.Image] | None,
    elements: list[Any] | None,
    crops: dict[str, bytes] | None,
    crops_dir: str | Path | None,
    assets_dir: str | Path | None,
) -> None:
    img_bytes = _resolve_image(
        block, page_images or [], elements, crops, crops_dir, assets_dir
    )
    if img_bytes:
        _add_image(doc, img_bytes, block.get("alt", ""))


def _add_image(doc: Document, img_bytes: bytes, alt: str = "") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(BytesIO(img_bytes), width=Inches(5.0))
    if alt:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(alt)
        cap_run.italic = True


def _add_table(doc: Document, block: dict) -> None:
    rows_data = block.get("rows", [])
    if not rows_data:
        return
    cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=cols)
    table.style = "Table Grid"
    for ri, row_data in enumerate(rows_data):
        for ci, cell_text in enumerate(row_data):
            if ci >= cols:
                break
            cell = table.cell(ri, ci)
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _ensure_rtl(paragraph)


def _add_formula(doc: Document, latex: str, display: bool) -> None:
    """Add a formula block — rendered as styled text with math font."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _ensure_rtl(p)

    run = p.add_run(latex)
    run.italic = True
    run.font.name = "Cambria Math"

    rpr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")  # ruff: ignore[non-lowercase-variable-in-function]
    rFonts.set(qn("w:ascii"), "Cambria Math")
    rFonts.set(qn("w:hAnsi"), "Cambria Math")
    rFonts.set(qn("w:cs"), "Cambria Math")
    rpr.append(rFonts)

    pPr = p._element.get_or_add_pPr()  # ruff: ignore[non-lowercase-variable-in-function]
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)


def _ensure_rtl(p: object) -> None:
    pPr = p._element.get_or_add_pPr()  # ruff: ignore[non-lowercase-variable-in-function]
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        el = OxmlElement("w:bidi")
        pPr.append(el)


# ── Markdown parsing ────────────────────────────────────────────────────────


def _try_append_structure_line(blocks: list[dict], stripped: str) -> bool:
    """Handle separator/heading/image/table lines; return True if consumed."""
    if stripped.startswith("<!-- page:") or stripped == "---":
        blocks.append({"kind": "separator"})
        return True
    heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
    if heading_match:
        blocks.append({
            "kind": "heading",
            "level": len(heading_match.group(1)),
            "text": heading_match.group(2),
        })
        return True
    img_match = re.match(r"!\[(.+?)\]\((.+?)\)", stripped)
    if img_match:
        blocks.append({
            "kind": "image",
            "alt": img_match.group(1),
            "src": img_match.group(2),
        })
        return True
    table_match = re.match(r"^\|(.+)\|$", stripped)
    if table_match:
        cells = [c.strip() for c in table_match.group(1).split("|")]
        if blocks and blocks[-1].get("kind") == "table":
            blocks[-1].setdefault("rows", []).append(cells)
        else:
            blocks.append({"kind": "table", "rows": [cells]})
        return True
    return bool(re.match(r"^\|[-:| ]+\|$", stripped))


def _append_markdown_line(blocks: list[dict], stripped: str) -> None:
    """Classify one non-empty markdown line into *blocks*."""
    if _try_append_structure_line(blocks, stripped):
        return
    formula_match = re.match(r"^\$\$(.+)\$\$$", stripped)
    if formula_match:
        blocks.append({
            "kind": "formula",
            "latex": formula_match.group(1),
            "display": True,
        })
        return
    inline_formula = re.match(r"^\$(.+)\$$", stripped)
    if inline_formula:
        blocks.append({
            "kind": "formula",
            "latex": inline_formula.group(1),
            "display": False,
        })
        return
    if stripped.startswith("> "):
        blocks.append({"kind": "blockquote", "text": stripped[2:]})
        return
    if re.match(r"^[\-\*]\s", stripped):
        blocks.append({
            "kind": "list_item",
            "text": re.sub(r"^[\-\*]\s+", "", stripped),
        })
        return
    if re.match(r"^\d+[.\)]\s", stripped):
        blocks.append({
            "kind": "list_item",
            "text": re.sub(r"^\d+[.\)]\s+", "", stripped),
        })
        return
    blocks.append({"kind": "paragraph", "text": stripped})


def _parse_markdown_blocks(md: str) -> list[dict]:
    blocks: list[dict] = []
    for line in md.split("\n"):
        stripped = line.strip()
        if stripped:
            _append_markdown_line(blocks, stripped)
    return blocks


def _bytes_from_src(src: str, assets_dir: str | Path | None) -> bytes | None:
    if not src or src == "#" or src.startswith("asset:"):
        return None
    path = Path(src)
    if path.exists():
        return path.read_bytes()
    if assets_dir:
        alt_path = Path(assets_dir) / path.name
        if alt_path.exists():
            return alt_path.read_bytes()
    return None


def _bytes_from_crop_cache(
    elements: list[Any] | None, crops: dict[str, bytes] | None
) -> bytes | None:
    if not (crops and elements):
        return None
    from .layout_detector import (
        ElementType as ET,  # ruff: ignore[camelcase-imported-as-acronym]
    )

    for elem in elements:
        if elem.type in (ET.figure, ET.chart):
            crop_bytes = crops.get(elem.element_id)
            if crop_bytes:
                return crop_bytes
    return None


def _bytes_from_page_crop(
    page_images: list[PILImage.Image], elements: list[Any] | None
) -> bytes | None:
    if not (page_images and elements):
        return None
    from .layout_detector import (
        ElementType as ET,  # ruff: ignore[camelcase-imported-as-acronym]
    )

    for elem in elements:
        if elem.type not in (ET.figure, ET.chart):
            continue
        idx = min(elem.page_number - 1, len(page_images) - 1)
        img = page_images[idx]
        x1 = max(0, int(elem.x1) - 5)
        y1 = max(0, int(elem.y1) - 5)
        x2 = min(img.width, int(elem.x2) + 5)
        y2 = min(img.height, int(elem.y2) + 5)
        if x2 > x1 and y2 > y1:
            crop = img.crop((x1, y1, x2, y2))
            buf = BytesIO()
            crop.save(buf, format="PNG")
            buf.seek(0)
            return buf.read()
    return None


def _resolve_image(
    block: dict,
    page_images: list[PILImage.Image],
    elements: list[Any] | None,
    crops: dict[str, bytes] | None,
    crops_dir: str | Path | None,
    assets_dir: str | Path | None,
) -> bytes | None:
    from_src = _bytes_from_src(block.get("src", ""), assets_dir)
    if from_src is not None:
        return from_src
    from_crops = _bytes_from_crop_cache(elements, crops)
    if from_crops is not None:
        return from_crops
    from_page = _bytes_from_page_crop(page_images, elements)
    if from_page is not None:
        return from_page
    if crops_dir:
        for f in Path(crops_dir).iterdir():
            if f.suffix in (".png", ".jpg", ".jpeg"):
                return f.read_bytes()
    return None
