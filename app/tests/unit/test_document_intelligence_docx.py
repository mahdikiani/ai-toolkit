"""Unit tests for the Document Intelligence Word (DOCX) renderer."""

from io import BytesIO
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Emu, Inches
from PIL import Image

from apps.ocr.document_intelligence.ast import ASTNode, DocumentAST, PageAST
from apps.ocr.document_intelligence.latex_omml import LatexConversionError
from apps.ocr.document_intelligence.layout import LayoutType
from apps.ocr.document_intelligence.renderers.docx import (
    _resolve_image_width_in,
    render_docx,
)


def _open(buf: BytesIO) -> WordDocument:
    buf.seek(0)
    return WordDocument(buf)


def _jc_val(paragraph: object) -> str | None:
    """
    Raw <w:jc w:val=".."/>, since python-docx's .alignment property
    raises on logical values ("start"/"end") it doesn't have an enum
    member for."""
    from docx.oxml.ns import qn

    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return None
    jc = p_pr.find(qn("w:jc"))
    return jc.get(qn("w:val")) if jc is not None else None


@pytest.mark.document_intelligence
class TestFormulaRendersRealOmml:
    """Formulas must become real, editable OMML equation objects, not text."""

    def test_valid_formula_produces_omath_xml(self) -> None:
        doc_ast = DocumentAST(
            pages=[
                PageAST(
                    page_number=1,
                    nodes=[ASTNode(type=LayoutType.formula, latex=r"\frac{x^2}{y}")],
                )
            ]
        )

        buf = render_docx(doc_ast)
        doc = _open(buf)

        assert "oMath" in doc.element.xml

    def test_invalid_formula_falls_back_to_styled_text_without_crashing(self) -> None:
        doc_ast = DocumentAST(
            pages=[
                PageAST(
                    page_number=1,
                    nodes=[ASTNode(type=LayoutType.formula, latex="E=mc^2")],
                )
            ]
        )

        with patch(
            "apps.ocr.document_intelligence.renderers.docx.latex_to_omml",
            side_effect=LatexConversionError("boom"),
        ):
            buf = render_docx(doc_ast)  # must not raise
        doc = _open(buf)

        assert "oMath" not in doc.element.xml
        fallback_runs = [
            run for p in doc.paragraphs for run in p.runs if run.text == "E=mc^2"
        ]
        assert fallback_runs
        assert fallback_runs[0].font.name == "Cambria Math"


@pytest.mark.document_intelligence
class TestTableRendersRealWordTable:
    def test_rows_become_a_real_table_object(self) -> None:
        node = ASTNode(type=LayoutType.table, rows=[["A", "B"], ["1", "2"]])
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert table.cell(0, 0).text == "A"
        assert table.cell(1, 1).text == "2"

    def test_cell_merges_produce_a_real_merged_word_cell(self) -> None:
        """
        rowspan/colspan reconstructed from HTML must become an actual
        merged cell (cell.merge()), not two separate cells with duplicated
        or blank text."""
        node = ASTNode(
            type=LayoutType.table,
            rows=[["Header", ""], ["A", "B"]],
            cell_merges=[(0, 0, 0, 1)],
        )
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        table = doc.tables[0]
        assert table.cell(0, 0).text.strip() == "Header"
        assert table.cell(0, 1).text.strip() == "Header"
        assert table.cell(0, 0)._tc is table.cell(0, 1)._tc

    def test_repeat_header_row_sets_real_tbl_header(self) -> None:
        node = ASTNode(
            type=LayoutType.table,
            rows=[["A", "B"], ["1", "2"]],
            repeat_header_row=True,
        )
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        table = doc.tables[0]
        assert "w:tblHeader" in table.rows[0]._tr.xml
        assert "w:tblHeader" not in table.rows[1]._tr.xml


@pytest.mark.document_intelligence
class TestHeaderFooterUseRealWordSections:
    """Headers/footers must land in doc.sections[].header/footer, not the body."""

    def test_repeated_header_and_footer_promoted_to_section(self) -> None:
        pages = [
            PageAST(
                page_number=n,
                nodes=[
                    ASTNode(type=LayoutType.header, text="Company Confidential"),
                    ASTNode(type=LayoutType.paragraph, text=f"body {n}"),
                    ASTNode(type=LayoutType.footer, text="Page footer"),
                ],
            )
            for n in (1, 2)
        ]
        doc_ast = DocumentAST(pages=pages)

        doc = _open(render_docx(doc_ast))

        assert doc.sections[0].header.paragraphs[0].text == "Company Confidential"
        assert doc.sections[0].footer.paragraphs[0].text == "Page footer"
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Company Confidential" not in body_text
        assert "Page footer" not in body_text

    def test_single_page_header_not_promoted_but_still_rendered(self) -> None:
        """
        A header/footer seen on only one page can't be verified as a
        repeating banner -- it must not be promoted globally, but it also
        must not be silently dropped: it falls back to an inline paragraph."""
        node = ASTNode(type=LayoutType.header, text="One-off notice")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        header_paras = doc.sections[0].header.paragraphs
        assert header_paras == [] or header_paras[0].text == ""
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "One-off notice" in body_text

    def test_header_seen_once_across_many_pages_not_promoted(self) -> None:
        """
        Repetition must be a real cross-page majority -- a value seen on
        1 of 5 pages is noise, not a genuine repeating header."""
        pages = [
            PageAST(
                page_number=1,
                nodes=[ASTNode(type=LayoutType.header, text="Chapter 3 intro")],
            ),
        ] + [PageAST(page_number=n, nodes=[]) for n in range(2, 6)]
        doc_ast = DocumentAST(pages=pages)

        doc = _open(render_docx(doc_ast))

        header_paras = doc.sections[0].header.paragraphs
        assert header_paras == [] or header_paras[0].text == ""
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Chapter 3 intro" in body_text

    def test_two_footer_regions_both_render_as_separate_paragraphs(self) -> None:
        """
        Regression test: writing a second promoted region onto the same
        footer part must not overwrite the first one's paragraph -- a
        repeating footer text plus an independent page-number sequence are
        two distinct paragraphs, not one clobbering the other."""
        pages = [
            PageAST(
                page_number=n,
                nodes=[
                    ASTNode(type=LayoutType.footer, text="Confidential"),
                    ASTNode(type=LayoutType.page_number, text=str(n)),
                ],
            )
            for n in (1, 2, 3)
        ]
        doc_ast = DocumentAST(pages=pages)

        doc = _open(render_docx(doc_ast))

        footer_paras = doc.sections[0].footer.paragraphs
        assert len(footer_paras) == 2
        assert footer_paras[0].text == "Confidential"
        assert "PAGE" in footer_paras[1]._p.xml

    def test_two_header_slots_render_as_two_separate_paragraphs(self) -> None:
        node_top = ASTNode(type=LayoutType.header, text="ACME Corp", bbox=(0, 10, 200, 40))
        node_bottom = ASTNode(
            type=LayoutType.header, text="Confidential Draft", bbox=(0, 60, 200, 90)
        )
        pages = [
            PageAST(page_number=n, nodes=[node_top, node_bottom], page_height=1000)
            for n in (1, 2, 3)
        ]
        doc_ast = DocumentAST(pages=pages)

        doc = _open(render_docx(doc_ast))

        header_paras = doc.sections[0].header.paragraphs
        assert [p.text for p in header_paras] == ["ACME Corp", "Confidential Draft"]

    def test_different_first_page_writes_distinct_first_page_header(self) -> None:
        pages = [
            PageAST(page_number=1, nodes=[ASTNode(type=LayoutType.header, text="Title Page")]),
            *[
                PageAST(page_number=n, nodes=[ASTNode(type=LayoutType.header, text="Chapter")])
                for n in (2, 3, 4)
            ],
        ]
        doc_ast = DocumentAST(pages=pages)

        doc = _open(render_docx(doc_ast))

        section = doc.sections[0]
        assert section.different_first_page_header_footer is True
        assert section.first_page_header.paragraphs[0].text == "Title Page"
        assert section.header.paragraphs[0].text == "Chapter"
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Title Page" not in body_text
        assert "Chapter" not in body_text


@pytest.mark.document_intelligence
class TestListRendersBulletOrNumberStyle:
    def test_bulleted_item_uses_list_bullet_style(self) -> None:
        node = ASTNode(
            type=LayoutType.list,
            children=[ASTNode(type=LayoutType.list, text="item one", ordered=False)],
        )
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        matching = [p for p in doc.paragraphs if p.text == "item one"]
        assert matching
        assert matching[0].style.name == "List Bullet"

    def test_numbered_item_uses_list_number_style(self) -> None:
        node = ASTNode(
            type=LayoutType.list,
            children=[ASTNode(type=LayoutType.list, text="item one", ordered=True)],
        )
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        matching = [p for p in doc.paragraphs if p.text == "item one"]
        assert matching
        assert matching[0].style.name == "List Number"


@pytest.mark.document_intelligence
class TestPreviouslyDroppedNodeTypesAreNowRendered:
    """
    table_caption/table_footnote/figure_caption/unknown used to fall
    through the dispatch with no branch and were silently lost -- every one
    of these must now show up somewhere in the document body."""

    @pytest.mark.parametrize(
        "layout_type",
        [
            LayoutType.table_caption,
            LayoutType.table_footnote,
            LayoutType.figure_caption,
            LayoutType.unknown,
        ],
    )
    def test_node_text_appears_in_body(self, layout_type) -> None:
        node = ASTNode(type=layout_type, text="این نباید گم شود")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "این نباید گم شود" in body_text

    def test_unverifiable_page_number_falls_back_to_inline_text(self) -> None:
        """
        A page_number node that can't be verified as a real cross-page
        sequence (e.g. a single-page document) is not promoted to a PAGE
        field, but must still render somewhere rather than being dropped."""
        node = ASTNode(type=LayoutType.page_number, text="7")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "7" in body_text

    def test_verified_page_number_sequence_becomes_real_page_field(self) -> None:
        """
        A page_number node matching the actual page index on every page
        is a genuine sequence -- it must become a real PAGE field in the
        footer, never raw OCR'd digits in the body."""
        pages = [
            PageAST(
                page_number=n, nodes=[ASTNode(type=LayoutType.page_number, text=str(n))]
            )
            for n in (1, 2, 3)
        ]
        doc_ast = DocumentAST(pages=pages)

        doc = _open(render_docx(doc_ast))

        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "1" not in body_text
        assert "2" not in body_text
        assert "3" not in body_text
        assert "PAGE" in doc.sections[0].footer.paragraphs[0]._p.xml


@pytest.mark.document_intelligence
class TestRtlLayout:
    def test_paragraphs_are_marked_bidi(self) -> None:
        node = ASTNode(type=LayoutType.paragraph, text="سلام دنیا")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))
        body_paragraphs = [p for p in doc.paragraphs if p.text == "سلام دنیا"]

        assert body_paragraphs
        assert "w:bidi" in body_paragraphs[0]._p.xml

    def test_bidi_precedes_jc_in_schema_order(self) -> None:
        """
        Regression: <w:bidi/> was being appended to pPr *after* <w:jc/>
        was already present (alignment is set before _ensure_rtl runs),
        violating the OOXML CT_PPr element sequence (bidi must precede jc).
        LibreOffice/python-docx tolerate the misordering silently, but real
        Microsoft Word does not -- observed directly: a paragraph reporting
        WD_ALIGN_PARAGRAPH.RIGHT and rendering correctly in LibreOffice
        showed up left-aligned when opened in real Word."""
        from docx.oxml.ns import qn

        node = ASTNode(type=LayoutType.paragraph, text="سلام دنیا")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))
        p = next(p for p in doc.paragraphs if p.text == "سلام دنیا")
        p_pr = p._p.find(qn("w:pPr"))
        tags = [child.tag.split("}")[-1] for child in p_pr]

        assert tags.index("bidi") < tags.index("jc")

    def test_lang_is_never_a_direct_child_of_ppr(self) -> None:
        """
        Regression: <w:lang> was being appended directly onto a paragraph's
        <w:pPr> (to set the paragraph-mark's proofing language), but
        <w:lang> is only valid OOXML nested inside a <w:rPr> -- it is not a
        member of CT_PPr's own child sequence (EG_PPrBase) at all. Real
        Word can repair/drop a paragraph's formatting on open when it
        finds this; LibreOffice/python-docx read it leniently and don't
        surface the problem, which is how it went unnoticed. The correct
        place is <w:pPr><w:rPr><w:lang/></w:rPr></w:pPr>."""
        from docx.oxml.ns import qn

        node = ASTNode(type=LayoutType.paragraph, text="سلام دنیا")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))
        p = next(p for p in doc.paragraphs if p.text == "سلام دنیا")
        p_pr = p._p.find(qn("w:pPr"))
        direct_children = [child.tag.split("}")[-1] for child in p_pr]

        assert "lang" not in direct_children
        rpr = p_pr.find(qn("w:rPr"))
        assert rpr is not None
        assert rpr.find(qn("w:lang")) is not None

    def test_section_is_marked_bidi(self) -> None:
        """
        Regression: per-paragraph <w:bidi/> alone was not enough -- observed
        directly in Microsoft Word for Mac, a document with every paragraph
        correctly carrying <w:bidi/> and <w:jc w:val="right"/> still
        rendered pinned to the left margin, because the *section* itself
        (w:sectPr) was never marked right-to-left. Word additionally needs
        this section-level flag alongside the per-paragraph one."""
        from docx.oxml.ns import qn

        node = ASTNode(type=LayoutType.paragraph, text="سلام دنیا")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))
        sect_pr = doc.sections[0]._sectPr

        assert sect_pr.find(qn("w:bidi")) is not None

    def test_section_bidi_precedes_doc_grid_in_schema_order(self) -> None:

        node = ASTNode(type=LayoutType.paragraph, text="سلام دنیا")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))
        sect_pr = doc.sections[0]._sectPr
        tags = [child.tag.split("}")[-1] for child in sect_pr]

        assert tags.index("bidi") < tags.index("docGrid")

    def test_code_block_stays_left_aligned_in_rtl_document(self) -> None:
        """
        Source code is LTR content regardless of the surrounding RTL
        document -- without an explicit override it would silently inherit
        the Normal style's RIGHT alignment."""
        node = ASTNode(type=LayoutType.code, text="print('hi')")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        matching = [p for p in doc.paragraphs if p.text == "print('hi')"]
        assert matching
        assert matching[0].alignment == WD_ALIGN_PARAGRAPH.LEFT


_WML_XSD_CANDIDATES = [
    "/home/mahdi/.grok/skills/docx/scripts/office/schemas/ISO-IEC29500-4_2016/wml.xsd",
]
_WML_XSD_PATH = next((p for p in _WML_XSD_CANDIDATES if Path(p).is_file()), None)


@pytest.mark.document_intelligence
@pytest.mark.skipif(
    _WML_XSD_PATH is None, reason="ISO/IEC 29500-4 WordprocessingML XSD not available"
)
class TestOoxmlSchemaValidity:
    """
    Validate generated XML against the real ISO/IEC 29500-4 WordprocessingML
    schema (not just python-docx's lenient reader), since python-docx and
    LibreOffice both tolerate schema violations that real Microsoft Word
    can repair-on-open, silently dropping the very formatting under test
    elsewhere in this file (see e.g. test_lang_is_never_a_direct_child_of_ppr
    and the module-history: bidi wrongly nested in a style's rPr, a
    style's pPr getting an invalid nested rPr, <w:shd/> appended out of
    schema order in code blocks -- all found only by validating against
    this real schema, not by any amount of reading python-docx's output).
    """

    def test_full_document_is_schema_valid(self) -> None:
        from lxml import etree

        nodes = [
            ASTNode(type=LayoutType.title, text="عنوان سند"),
            ASTNode(type=LayoutType.heading, text="زیرعنوان یک", level=1),
            ASTNode(type=LayoutType.heading, text="زیرعنوان دو", level=2),
            ASTNode(type=LayoutType.heading, text="زیرعنوان سه", level=3),
            ASTNode(type=LayoutType.paragraph, text="این یک پاراگراف فارسی است."),
            ASTNode(type=LayoutType.list, text="آیتم یک\nآیتم دو", ordered=True),
            ASTNode(type=LayoutType.list, text="آیتم سه\nآیتم چهار", ordered=False),
            ASTNode(type=LayoutType.code, text="print(1)"),
            ASTNode(
                type=LayoutType.table,
                rows=[["A", "B"], ["1", "2"]],
                repeat_header_row=True,
            ),
            ASTNode(type=LayoutType.formula, latex=r"\frac{x^2}{y}"),
            ASTNode(type=LayoutType.header, text="Company Confidential"),
            ASTNode(type=LayoutType.footer, text="Page footer"),
            ASTNode(type=LayoutType.table_caption, text="caption"),
            ASTNode(type=LayoutType.reference, text="ref"),
            ASTNode(type=LayoutType.unknown, text="unknown text"),
        ]
        pages = [PageAST(page_number=n, nodes=nodes) for n in (1, 2)]
        doc_ast = DocumentAST(pages=pages)
        buf = render_docx(doc_ast)

        schema = etree.XMLSchema(etree.parse(_WML_XSD_PATH))
        import zipfile

        z = zipfile.ZipFile(buf)
        failures = []
        wml_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        for name in z.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            if name == "word/settings.xml":
                # Untouched python-docx template content (verified: identical
                # in a pristine Document(), before any of our code runs) --
                # its zoom/docId MCE fallback content isn't ours to fix and
                # isn't what this test is guarding against.
                continue
            tree = etree.fromstring(z.read(name))
            if not tree.tag.startswith(wml_ns):
                continue
            if schema.validate(tree):
                continue
            # mc:Ignorable (Markup Compatibility) is a real, universally
            # supported OOXML mechanism this single-schema validator
            # doesn't know about -- it appears unmodified in every
            # python-docx document (confirmed against a pristine
            # Document() template), so it's excluded as a known false
            # positive rather than a defect in our own generated content.
            real_errors = [e for e in schema.error_log if "Ignorable" not in str(e)]
            if real_errors:
                failures.append((name, [str(e) for e in real_errors]))
        assert not failures, failures


@pytest.mark.document_intelligence
class TestDocumentTitleMetadata:
    def test_title_set_on_core_properties(self) -> None:
        doc_ast = DocumentAST(title="My Report", pages=[])

        doc = _open(render_docx(doc_ast))

        assert doc.core_properties.title == "My Report"


@pytest.mark.document_intelligence
class TestFontDetection:
    """
    The renderer must use the source PDF's real fonts when available,
    not always fall back to the fixed Calibri/IRNazanin defaults."""

    def test_detected_fonts_applied_to_normal_style(self) -> None:
        node = ASTNode(type=LayoutType.paragraph, text="hello")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        with patch(
            "apps.ocr.pipeline.docx_renderer.detect_pdf_fonts",
            return_value={"cs": "Vazir", "latin": "Arial"},
        ):
            doc = _open(render_docx(doc_ast, pdf_data=b"%PDF-fake"))

        xml = doc.styles["Normal"].element.xml
        assert 'w:cs="Vazir"' in xml
        assert 'w:ascii="Arial"' in xml

    def test_no_pdf_data_uses_default_fonts(self) -> None:
        node = ASTNode(type=LayoutType.paragraph, text="hello")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        xml = doc.styles["Normal"].element.xml
        assert 'w:cs="IRNazanin"' in xml

    def test_font_detection_failure_falls_back_without_crashing(self) -> None:
        node = ASTNode(type=LayoutType.paragraph, text="hello")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        with patch(
            "apps.ocr.pipeline.docx_renderer.detect_pdf_fonts",
            side_effect=RuntimeError("boom"),
        ):
            doc = _open(render_docx(doc_ast, pdf_data=b"%PDF-fake"))  # must not raise

        assert 'w:cs="IRNazanin"' in doc.styles["Normal"].element.xml


@pytest.mark.document_intelligence
class TestImageSizing:
    """
    Images must be sized relative to their footprint on the source page,
    not always dropped in at a fixed 5 inches."""

    def test_small_bbox_produces_small_width(self) -> None:
        width = _resolve_image_width_in(
            ASTNode(type=LayoutType.figure, bbox=(0, 0, 200, 200)),
            page_width_px=1000,
            content_width_in=6.0,
        )
        assert width == pytest.approx(1.5)  # 20% of content width, but floored at MIN

    def test_full_width_bbox_fills_content_width(self) -> None:
        width = _resolve_image_width_in(
            ASTNode(type=LayoutType.figure, bbox=(0, 0, 950, 200)),
            page_width_px=1000,
            content_width_in=6.0,
        )
        assert width == pytest.approx(5.7)

    def test_missing_bbox_falls_back_to_capped_default(self) -> None:
        width = _resolve_image_width_in(
            ASTNode(type=LayoutType.figure), page_width_px=1000, content_width_in=4.0
        )
        assert width == pytest.approx(4.0)  # capped to content width, not fixed 5"

    def test_rendered_image_width_matches_bbox_ratio(self, tmp_path) -> None:
        img_path = tmp_path / "fig.png"
        Image.new("RGB", (50, 50), "white").save(img_path)

        node = ASTNode(
            type=LayoutType.figure,
            asset_path=str(img_path),
            bbox=(0, 0, 500, 300),  # 50% of a 1000px-wide page
        )
        page = PageAST(
            page_number=1, nodes=[node], page_width=1000, page_height=1400, page_dpi=100
        )
        doc_ast = DocumentAST(pages=[page])

        doc = _open(render_docx(doc_ast))

        pic = doc.inline_shapes[0]
        # content width = 10in page (1000px/100dpi) minus 2*0.8in margin = 8.4in;
        # 50% of that = 4.2in
        assert pic.width == Emu(Inches(4.2))


@pytest.mark.document_intelligence
class TestImageDescriptionBecomesAltText:
    """
    A VLM-generated image/chart description is accessibility metadata.

    It must land in the image's real Word "Alt Text" field
    (<wp:docPr descr="..."/>), not as a second, visible body paragraph
    duplicating what a sighted reader already sees in the picture.
    """

    def test_figure_description_sets_alt_text_not_visible_paragraph(
        self, tmp_path
    ) -> None:
        img_path = tmp_path / "fig.png"
        Image.new("RGB", (50, 50), "white").save(img_path)
        node = ASTNode(
            type=LayoutType.figure,
            asset_path=str(img_path),
            description="A scatter plot showing two distinct groups of data.",
        )
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        from docx.oxml.ns import qn

        doc_pr = doc.element.body.find(f".//{qn('wp:docPr')}")
        assert doc_pr is not None
        assert doc_pr.get("descr") == (
            "A scatter plot showing two distinct groups of data."
        )
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "scatter plot" not in body_text

    def test_chart_description_sets_alt_text_not_visible_paragraph(
        self, tmp_path
    ) -> None:
        img_path = tmp_path / "chart.png"
        Image.new("RGB", (50, 50), "white").save(img_path)
        node = ASTNode(
            type=LayoutType.chart,
            asset_path=str(img_path),
            caption="Figure 1",
            description="Bar chart comparing quarterly revenue.",
        )
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Bar chart comparing" not in body_text
        assert "Figure 1" not in body_text

    def test_caption_is_never_shown_visibly_even_when_short(self, tmp_path) -> None:
        """
        Regression: a short VLM caption ("Young Scholars Club logo") is
        just as much an invented label as a long description -- there's
        no reliable way to tell it apart from a real printed caption
        (those are captured as their own separate text element by the
        layout detector, not through this VLM call at all), so it must
        never render as visible text regardless of length."""
        img_path = tmp_path / "logo.png"
        Image.new("RGB", (50, 50), "white").save(img_path)
        node = ASTNode(
            type=LayoutType.figure,
            asset_path=str(img_path),
            caption="Young Scholars Club logo",
        )
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Young Scholars Club logo" not in body_text
        from docx.oxml.ns import qn

        doc_pr = doc.element.body.find(f".//{qn('wp:docPr')}")
        assert doc_pr is not None
        assert doc_pr.get("descr") == "Young Scholars Club logo"


@pytest.mark.document_intelligence
class TestInlineMarkdownRendering:
    """
    Bold/italic/code spans and links inside paragraph/list/table text
    must become real per-span runs (or a real hyperlink), not raw
    ``**markers**`` dumped into a single plain run."""

    def test_bold_span_becomes_a_bold_run(self) -> None:
        node = ASTNode(type=LayoutType.paragraph, text="normal **bold** normal")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        bold_runs = [r for p in doc.paragraphs for r in p.runs if r.text == "bold"]
        assert bold_runs
        assert bold_runs[0].bold is True

    def test_code_span_uses_monospace_font(self) -> None:
        node = ASTNode(type=LayoutType.paragraph, text="run `code` here")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        code_runs = [r for p in doc.paragraphs for r in p.runs if r.text == "code"]
        assert code_runs
        assert code_runs[0].font.name == "Courier New"

    def test_link_becomes_a_real_hyperlink(self) -> None:
        node = ASTNode(
            type=LayoutType.paragraph, text="see [our docs](https://example.com) here"
        )
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        matching = [p for p in doc.paragraphs if "our docs" in p.text]
        assert matching
        p = matching[0]
        assert "w:hyperlink" in p._p.xml
        assert "our docs" in p._p.xml
        # A real relationship must exist, not just a styled run.
        rels = [r.target_ref for r in doc.part.rels.values() if r.is_external]
        assert "https://example.com" in rels

    def test_list_item_supports_bold_span(self) -> None:
        node = ASTNode(
            type=LayoutType.list,
            children=[ASTNode(type=LayoutType.list, text="**important** item")],
        )
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        bold_runs = [r for p in doc.paragraphs for r in p.runs if r.text == "important"]
        assert bold_runs
        assert bold_runs[0].bold is True

    def test_table_cell_supports_bold_span(self) -> None:
        node = ASTNode(type=LayoutType.table, rows=[["**A**", "B"]])
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        table = doc.tables[0]
        cell_runs = [r for p in table.cell(0, 0).paragraphs for r in p.runs]
        assert any(r.bold for r in cell_runs)

    def test_inline_math_becomes_a_real_omath_object(self) -> None:
        """
        Mid-sentence LaTeX ($p_i$) must become a real inline m:oMath,
        not raw $-delimited text left inside the RTL paragraph (which
        produces bidi word-reordering artifacts -- observed directly on a
        real OCR'd document)."""
        node = ASTNode(type=LayoutType.paragraph, text="احتمال $p_i$ برابر است")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        matching = [p for p in doc.paragraphs if "احتمال" in p.text]
        assert matching
        assert "m:oMath" in matching[0]._p.xml
        assert "$" not in matching[0]._p.xml

    def test_invalid_inline_math_falls_back_to_styled_text(self) -> None:
        node = ASTNode(type=LayoutType.paragraph, text="بد $\\notarealcommand{$ متن")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))  # must not raise

        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "بد" in body_text


@pytest.mark.document_intelligence
class TestParagraphAlignment:
    def test_centered_narrow_bbox_gets_center_alignment(self) -> None:
        node = ASTNode(
            type=LayoutType.paragraph, text="centered", bbox=(300, 0, 700, 30)
        )
        page = PageAST(page_number=1, nodes=[node], page_width=1000, page_height=1400)
        doc_ast = DocumentAST(pages=[page])

        doc = _open(render_docx(doc_ast))

        matching = [p for p in doc.paragraphs if p.text == "centered"]
        assert matching
        assert matching[0].alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_off_center_bbox_stays_right_aligned(self) -> None:
        node = ASTNode(
            type=LayoutType.paragraph, text="right side", bbox=(700, 0, 950, 30)
        )
        page = PageAST(page_number=1, nodes=[node], page_width=1000, page_height=1400)
        doc_ast = DocumentAST(pages=[page])

        doc = _open(render_docx(doc_ast))

        matching = [p for p in doc.paragraphs if p.text == "right side"]
        assert matching
        assert _jc_val(matching[0]) == "start"

    def test_missing_bbox_defaults_to_right_aligned(self) -> None:
        node = ASTNode(type=LayoutType.paragraph, text="no bbox")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        doc = _open(render_docx(doc_ast))

        matching = [p for p in doc.paragraphs if p.text == "no bbox"]
        assert matching
        assert _jc_val(matching[0]) == "start"


@pytest.mark.document_intelligence
class TestPageSize:
    def test_page_size_matches_source_page_dimensions(self) -> None:
        # 1000x1400 px at 100 dpi -> 10in x 14in
        page = PageAST(
            page_number=1, nodes=[], page_width=1000, page_height=1400, page_dpi=100
        )
        doc_ast = DocumentAST(pages=[page])

        doc = _open(render_docx(doc_ast))

        section = doc.sections[0]
        assert section.page_width == Emu(Inches(10.0))
        assert section.page_height == Emu(Inches(14.0))

    def test_missing_page_dims_fall_back_to_a4(self) -> None:
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[])])

        doc = _open(render_docx(doc_ast))

        section = doc.sections[0]
        # Section sizing round-trips through twips (1/20 pt) internally, so
        # a non-round inch value like 8.27" can be off by a few EMU.
        assert section.page_width.inches == pytest.approx(8.27, abs=0.01)


@pytest.mark.document_intelligence
class TestMultiSection:
    def test_landscape_page_gets_its_own_real_section(self) -> None:
        """
        A page-size change partway through the document must produce a
        real new Word section (its own sectPr with its own orientation),
        not just a page break at the first page's dimensions."""
        pages = [
            PageAST(
                page_number=1,
                nodes=[ASTNode(type=LayoutType.paragraph, text="portrait page")],
                page_width=1000,
                page_height=1400,
                page_dpi=100,
            ),
            PageAST(
                page_number=2,
                nodes=[ASTNode(type=LayoutType.paragraph, text="landscape page")],
                page_width=1400,
                page_height=1000,
                page_dpi=100,
            ),
        ]
        doc_ast = DocumentAST(pages=pages)

        doc = _open(render_docx(doc_ast))

        assert len(doc.sections) == 2
        assert doc.sections[0].page_width < doc.sections[0].page_height  # portrait
        assert doc.sections[1].page_width > doc.sections[1].page_height  # landscape
        assert doc.sections[1].page_width == Emu(Inches(14.0))
        assert doc.sections[1].page_height == Emu(Inches(10.0))

    def test_uniform_document_stays_a_single_section(self) -> None:
        pages = [
            PageAST(
                page_number=n,
                nodes=[ASTNode(type=LayoutType.paragraph, text=f"page {n}")],
                page_width=1000,
                page_height=1400,
                page_dpi=100,
            )
            for n in (1, 2, 3)
        ]
        doc_ast = DocumentAST(pages=pages)

        doc = _open(render_docx(doc_ast))

        assert len(doc.sections) == 1

    def test_content_width_adapts_to_the_new_section(self) -> None:
        """
        Content sized/aligned against page width must use the section
        it's actually rendered in, not the first section's width."""
        landscape_node = ASTNode(
            type=LayoutType.figure, asset_path="", bbox=(0, 0, 700, 200)
        )
        pages = [
            PageAST(
                page_number=1,
                nodes=[ASTNode(type=LayoutType.paragraph, text="p1")],
                page_width=1000,
                page_height=1400,
                page_dpi=100,
            ),
            PageAST(
                page_number=2,
                nodes=[landscape_node],
                page_width=1400,
                page_height=1000,
                page_dpi=100,
            ),
        ]
        doc_ast = DocumentAST(pages=pages)

        # Must not raise -- content on the second (landscape) section is
        # rendered against that section's own content width.
        doc = _open(render_docx(doc_ast))
        assert len(doc.sections) == 2

    def test_multi_column_page_gets_real_word_columns(self) -> None:
        pages = [
            PageAST(
                page_number=1,
                nodes=[ASTNode(type=LayoutType.paragraph, text="two column text")],
                page_width=1000,
                page_height=1400,
                page_dpi=100,
                column_count=2,
            )
        ]
        doc_ast = DocumentAST(pages=pages)

        doc = _open(render_docx(doc_ast))

        assert 'w:num="2"' in doc.sections[0]._sectPr.xml

    def test_single_column_page_keeps_default_one_column(self) -> None:
        pages = [
            PageAST(
                page_number=1,
                nodes=[ASTNode(type=LayoutType.paragraph, text="normal text")],
                page_width=1000,
                page_height=1400,
                page_dpi=100,
            )
        ]
        doc_ast = DocumentAST(pages=pages)

        doc = _open(render_docx(doc_ast))

        assert 'w:num="1"' in doc.sections[0]._sectPr.xml

    def test_column_count_change_produces_its_own_section(self) -> None:
        pages = [
            PageAST(
                page_number=1,
                nodes=[ASTNode(type=LayoutType.paragraph, text="single column")],
                page_width=1000,
                page_height=1400,
                page_dpi=100,
                column_count=1,
            ),
            PageAST(
                page_number=2,
                nodes=[ASTNode(type=LayoutType.paragraph, text="two column")],
                page_width=1000,
                page_height=1400,
                page_dpi=100,
                column_count=2,
            ),
        ]
        doc_ast = DocumentAST(pages=pages)

        doc = _open(render_docx(doc_ast))

        assert len(doc.sections) == 2
        assert 'w:num="1"' in doc.sections[0]._sectPr.xml
        assert 'w:num="2"' in doc.sections[1]._sectPr.xml
