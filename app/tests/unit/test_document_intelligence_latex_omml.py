"""Unit tests for LaTeX -> OMML conversion."""

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from apps.ocr.document_intelligence.latex_omml import (
    LatexConversionError,
    latex_to_omml,
)


@pytest.mark.unit
class TestLatexToOmml:
    def test_empty_latex_raises(self) -> None:
        with pytest.raises(LatexConversionError):
            latex_to_omml("")

    def test_fraction_produces_omml_fraction_element(self) -> None:
        omath = latex_to_omml(r"\frac{x^2}{y}")
        assert "<m:f>" in omath
        assert "<m:num>" in omath
        assert "<m:den>" in omath

    def test_superscript_produces_ssup(self) -> None:
        omath = latex_to_omml("x^2")
        assert "m:sSup" in omath

    def test_sum_with_limits_produces_subsup(self) -> None:
        omath = latex_to_omml(r"\sum_{i=1}^n x_i")
        assert "m:sSubSup" in omath

    def test_strips_dollar_delimiters(self) -> None:
        assert latex_to_omml("$x$") == latex_to_omml("x")

    @pytest.mark.parametrize(
        "tex",
        [r"\frac{x^2}{y}", "E=mc^2", r"\sum_{i=1}^n x_i", r"\sqrt{a_i}", r"\int_0^1 x dx"],
    )
    def test_round_trips_through_a_real_docx(self, tex: str) -> None:
        """The produced XML must be valid enough for python-docx to embed and re-save."""
        omath = latex_to_omml(tex)
        xml_str = f'<m:oMathPara {nsdecls("m")}>{omath}</m:oMathPara>'
        element = parse_xml(xml_str)

        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph._p.append(element)

        # Must save and re-open without raising.
        from io import BytesIO

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        Document(buf)
