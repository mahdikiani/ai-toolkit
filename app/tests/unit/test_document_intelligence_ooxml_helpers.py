"""
Unit tests for raw Open XML helpers (renderers/ooxml_helpers.py) --
Word features python-docx has no API for: field codes, real hyperlinks,
TOC fields."""

from io import BytesIO

import pytest
from docx import Document as WordDocument

from apps.ocr.document_intelligence.renderers.ooxml_helpers import (
    add_field_run,
    add_hyperlink_run,
    add_toc_field,
)


def _roundtrip(doc: WordDocument) -> WordDocument:
    """
    Save and reopen -- proves the XML is actually valid enough for
    python-docx (and by extension Word) to parse back, not just well-formed
    at construction time."""
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return WordDocument(buf)


@pytest.mark.document_intelligence
class TestAddFieldRun:
    def test_produces_real_field_code_xml(self) -> None:
        doc = WordDocument()
        p = doc.add_paragraph()

        add_field_run(p, "PAGE", cached_text="1")

        assert "<w:instrText" in p._p.xml
        assert "PAGE" in p._p.xml
        assert '<w:fldChar w:fldCharType="begin"/>' in p._p.xml
        assert '<w:fldChar w:fldCharType="end"/>' in p._p.xml

    def test_document_round_trips_through_save_and_reopen(self) -> None:
        doc = WordDocument()
        p = doc.add_paragraph()
        add_field_run(p, "NUMPAGES", cached_text="3")

        reopened = _roundtrip(doc)

        assert "NUMPAGES" in reopened.paragraphs[0]._p.xml

    def test_no_cached_text_still_produces_valid_field(self) -> None:
        doc = WordDocument()
        p = doc.add_paragraph()

        add_field_run(p, "PAGE")  # no cached_text

        reopened = _roundtrip(doc)
        assert "PAGE" in reopened.paragraphs[0]._p.xml


@pytest.mark.document_intelligence
class TestAddHyperlinkRun:
    def test_produces_real_hyperlink_element(self) -> None:
        doc = WordDocument()
        p = doc.add_paragraph()

        add_hyperlink_run(p, "click here", "https://example.com")

        assert "<w:hyperlink" in p._p.xml
        assert "click here" in p._p.xml

    def test_creates_a_real_external_relationship(self) -> None:
        doc = WordDocument()
        p = doc.add_paragraph()

        add_hyperlink_run(p, "click here", "https://example.com/page")

        rels = [r for r in doc.part.rels.values() if r.is_external]
        assert any(r.target_ref == "https://example.com/page" for r in rels)

    def test_run_carries_fa_ir_language(self) -> None:
        """
        Hyperlink runs live inside <w:hyperlink>, outside Paragraph.runs
        -- they must carry RTL language directly since the usual
        _ensure_rtl pass in docx.py can't reach them."""
        doc = WordDocument()
        p = doc.add_paragraph()

        add_hyperlink_run(p, "پیوند", "https://example.com")

        assert 'w:val="fa-IR"' in p._p.xml

    def test_document_round_trips_through_save_and_reopen(self) -> None:
        doc = WordDocument()
        p = doc.add_paragraph()
        add_hyperlink_run(p, "click here", "https://example.com")

        reopened = _roundtrip(doc)

        assert "w:hyperlink" in reopened.paragraphs[0]._p.xml


@pytest.mark.document_intelligence
class TestAddTocField:
    def test_produces_real_toc_instruction(self) -> None:
        doc = WordDocument()
        p = doc.add_paragraph()

        add_toc_field(p)

        assert "TOC" in p._p.xml
        assert 'w:fldCharType="begin"' in p._p.xml

    def test_document_round_trips_through_save_and_reopen(self) -> None:
        doc = WordDocument()
        p = doc.add_paragraph()
        add_toc_field(p)

        reopened = _roundtrip(doc)

        assert "TOC" in reopened.paragraphs[0]._p.xml
