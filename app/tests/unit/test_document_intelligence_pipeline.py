"""
Unit tests for the Document Intelligence orchestrator (full pipeline wiring).

Layout detection (PaddleOCR models) and the VLM HTTP call are mocked out —
this exercises the wiring between Loader -> Layout -> Elements -> Reading
Order -> AST -> Renderers -> Assets without needing GPU/network access.
"""

import asyncio
import time
from io import BytesIO
from unittest.mock import AsyncMock, patch

import pytest
from docx import Document as WordDocument
from PIL import Image

from apps.ocr.document_intelligence.layout import LayoutElement, LayoutType
from apps.ocr.document_intelligence.pipeline import (
    DocumentIntelligencePipeline,
    PipelineResult,
    summarize_stats,
)
from apps.ocr.document_intelligence.qa import run_docx_qa


def _fake_detect_page(_self, image: Image.Image, page) -> list[LayoutElement]:
    w, _h = image.size

    def elem(
        suffix: str, elem_type: LayoutType, box: tuple[int, int, int, int]
    ) -> LayoutElement:
        return LayoutElement(
            id=f"{page.id}_{suffix}",
            page_id=page.id,
            page_number=page.page_number,
            type=elem_type,
            bbox=box,
            padded_bbox=box,
            confidence=0.9,
        )

    return [
        elem("title", LayoutType.title, (10, 10, w - 10, 60)),
        elem("para", LayoutType.paragraph, (10, 80, w - 10, 200)),
        elem("table", LayoutType.table, (10, 220, w - 10, 320)),
        elem("formula", LayoutType.formula, (10, 340, w - 10, 380)),
        elem("figure", LayoutType.figure, (10, 400, w - 10, 500)),
    ]


async def _fake_vlm_call(
    _self,
    _crop,
    system_prompt: str,
    _user_prompt: str,
    response_format=None,
    max_tokens: int = 1024,
) -> str:
    _self._last_tokens = 7
    if "table extraction engine" in system_prompt:
        return "<table><tr><td>A</td><td>B</td></tr></table>"
    if "LaTeX source of the formula" in system_prompt:
        return r"\frac{x^2}{y}"
    if "Respond in exactly two lines" in system_prompt:
        return "caption: a test figure\ndescription: shows a test"
    return "سلام دنیا؛ این یک متن آزمایشی است."


def _blank_image(width: int = 600, height: int = 800) -> BytesIO:
    img = Image.new("RGB", (width, height), "white")
    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


@pytest.mark.document_intelligence
class TestDocumentIntelligencePipeline:
    """End-to-end wiring test with layout/VLM mocked."""

    async def _run(self, mode: str = "semantic") -> PipelineResult:
        with (
            patch(
                "apps.ocr.document_intelligence.layout.LayoutDetector.detect_page",
                _fake_detect_page,
            ),
            patch(
                "apps.ocr.document_intelligence.elements.ElementProcessor._vlm_call",
                _fake_vlm_call,
            ),
        ):
            pipeline = DocumentIntelligencePipeline(dpi=100)
            try:
                return await pipeline.process(_blank_image(), "test.png", mode=mode)
            finally:
                pipeline.cleanup()

    async def test_produces_markdown_with_all_element_types(self) -> None:
        result = await self._run()

        assert "# " in result.markdown  # title
        assert "| A | B |" in result.markdown  # table
        assert "$$" in result.markdown  # formula
        assert "assets/figure_" in result.markdown  # figure asset link
        assert "سلام دنیا" in result.markdown  # OCR'd paragraph text

    async def test_default_mode_produces_semantic_docx_with_real_table_and_omath(
        self,
    ) -> None:
        """
        Default mode ("semantic") must go through the flow-based
        renderer (docx.py): the table is a real top-level body table
        (doc.tables sees it directly), not boxed inside w:txbxContent."""
        result = await self._run()

        doc = WordDocument(BytesIO(result.docx_bytes))
        assert "w:txbxContent" not in doc.element.xml
        assert len(doc.tables) == 1
        assert "oMath" in doc.element.xml

    async def test_visual_mode_produces_absolute_layout_docx(self) -> None:
        """
        mode="visual" is the optional, non-default absolute-layout
        renderer — tables there are boxed per-element, not top-level."""
        result = await self._run(mode="visual")

        doc = WordDocument(BytesIO(result.docx_bytes))
        assert "w:txbxContent" in doc.element.xml
        assert "<w:tbl>" in doc.element.xml
        assert "oMath" in doc.element.xml

    async def test_saves_visual_element_as_asset(self) -> None:
        result = await self._run()

        assert len(result.assets) == 1
        assert result.assets[0].type == "figure"

    async def test_stats_cover_every_page_and_element(self) -> None:
        result = await self._run()

        assert len(result.stats.pages) == 1
        page_stats = result.stats.pages[0]
        assert page_stats.layout_time >= 0
        assert page_stats.vlm_time >= 0
        assert len(page_stats.elements) == 5

        summary = summarize_stats(result.stats)
        assert summary["pages"][0]["element_count"] == 5
        assert "elements" not in summary["pages"][0]  # debug detail hidden by default

        summary_debug = summarize_stats(result.stats, include_elements=True)
        assert "elements" in summary_debug["pages"][0]

    async def test_writes_local_output_files(self) -> None:
        result = await self._run()

        assert (result.output_dir / "document.md").exists()
        assert (result.output_dir / "document.docx").exists()
        assert any((result.output_dir / "assets").iterdir())

    async def test_default_mode_output_passes_the_qa_gate(self) -> None:
        """
        CI gate: the default (semantic) pipeline output must clear every
        QA check — no text boxes, nothing dropped, real tables/OMML,
        correct RTL, header/footer handled correctly."""
        result = await self._run()

        report = run_docx_qa(result.document_ast, result.docx_bytes, mode="semantic")

        assert report.passed, report.failures()


@pytest.mark.document_intelligence
class TestPageFaultIsolation:
    """
    Regression check.

    A single page failing after exhausted VLM retries used to abort
    the entire document, discarding every other already-processed page
    in what can be a multi-hour, thousand-call job for a large PDF.
    """

    async def test_one_failed_page_does_not_abort_the_document(self) -> None:
        from apps.ocr.document_intelligence.loader import Document, Page

        pages = [
            Page(
                id=f"p{i}",
                page_number=i,
                image=Image.new("RGB", (600, 800), "white"),
                width=600,
                height=800,
            )
            for i in range(1, 4)
        ]
        fake_document = Document(
            id="doc1",
            filename="test.pdf",
            pages=pages,
            pages_count=3,
            created_at="now",
        )

        pipeline = DocumentIntelligencePipeline(dpi=100)
        try:
            original_process_page = pipeline._process_page

            async def flaky_process_page(page):
                if page.page_number == 2:
                    error = RuntimeError("VLM exhausted retries")
                    raise error
                return await original_process_page(page)

            with (
                patch(
                    "apps.ocr.document_intelligence.pipeline.load_document",
                    return_value=fake_document,
                ),
                patch(
                    "apps.ocr.document_intelligence.layout.LayoutDetector.detect_page",
                    _fake_detect_page,
                ),
                patch(
                    "apps.ocr.document_intelligence.elements.ElementProcessor._vlm_call",
                    _fake_vlm_call,
                ),
                patch.object(pipeline, "_process_page", flaky_process_page),
            ):
                result = await pipeline.process(_blank_image(), "test.pdf")
        finally:
            pipeline.cleanup()

        assert len(result.stats.pages) == 3
        failed = [p for p in result.stats.pages if p.failed]
        assert [p.page_number for p in failed] == [2]
        succeeded = {p.page_number for p in result.stats.pages if not p.failed}
        assert succeeded == {1, 3}

        summary = summarize_stats(result.stats)
        assert summary["failed_pages"] == [2]


@pytest.mark.document_intelligence
class TestProgressCallback:
    """``on_page_done`` lets callers report progress on long documents."""

    async def test_called_after_every_page_including_failed_ones(self) -> None:
        from apps.ocr.document_intelligence.loader import Document, Page

        pages = [
            Page(
                id=f"p{i}",
                page_number=i,
                image=Image.new("RGB", (600, 800), "white"),
                width=600,
                height=800,
            )
            for i in range(1, 4)
        ]
        fake_document = Document(
            id="doc1",
            filename="test.pdf",
            pages=pages,
            pages_count=3,
            created_at="now",
        )

        pipeline = DocumentIntelligencePipeline(dpi=100)
        progress_calls: list[tuple[int, int]] = []

        async def on_page_done(completed: int, total: int) -> None:
            progress_calls.append((completed, total))

        try:
            original_process_page = pipeline._process_page

            async def flaky_process_page(page):
                if page.page_number == 2:
                    error = RuntimeError("VLM exhausted retries")
                    raise error
                return await original_process_page(page)

            with (
                patch(
                    "apps.ocr.document_intelligence.pipeline.load_document",
                    return_value=fake_document,
                ),
                patch(
                    "apps.ocr.document_intelligence.layout.LayoutDetector.detect_page",
                    _fake_detect_page,
                ),
                patch(
                    "apps.ocr.document_intelligence.elements.ElementProcessor._vlm_call",
                    _fake_vlm_call,
                ),
                patch.object(pipeline, "_process_page", flaky_process_page),
            ):
                await pipeline.process(
                    _blank_image(), "test.pdf", on_page_done=on_page_done
                )
        finally:
            pipeline.cleanup()

        assert progress_calls == [(1, 3), (2, 3), (3, 3)]

    async def test_broken_callback_does_not_abort_the_document(self) -> None:
        async def broken_on_page_done(_completed: int, _total: int) -> None:
            error = RuntimeError("webhook down")
            raise error

        with (
            patch(
                "apps.ocr.document_intelligence.layout.LayoutDetector.detect_page",
                _fake_detect_page,
            ),
            patch(
                "apps.ocr.document_intelligence.elements.ElementProcessor._vlm_call",
                _fake_vlm_call,
            ),
        ):
            pipeline = DocumentIntelligencePipeline(dpi=100)
            try:
                result = await pipeline.process(
                    _blank_image(), "test.png", on_page_done=broken_on_page_done
                )
            finally:
                pipeline.cleanup()

        assert "سلام دنیا" in result.markdown


@pytest.mark.document_intelligence
class TestCheckpointResume:
    """A page found in Redis checkpoints must not go through the VLM again."""

    async def test_checkpointed_page_is_reused_instead_of_reprocessed(self) -> None:
        from apps.ocr.document_intelligence.loader import Document, Page

        pages = [
            Page(
                id=f"p{i}",
                page_number=i,
                image=Image.new("RGB", (600, 800), "white"),
                width=600,
                height=800,
            )
            for i in range(1, 3)
        ]
        fake_document = Document(
            id="doc1", filename="test.pdf", pages=pages, pages_count=2, created_at="now"
        )
        checkpoints = {
            1: {
                "page_ast": {
                    "page_number": 1,
                    "nodes": [
                        {
                            "type": "paragraph",
                            "text": "cached page one text",
                            "children": [],
                        }
                    ],
                },
                "page_stat": {
                    "page_number": 1,
                    "layout_time": 0.0,
                    "vlm_time": 0.0,
                },
            }
        }

        pipeline = DocumentIntelligencePipeline(dpi=100)
        try:
            original_process_page = pipeline._process_page
            process_page_calls: list[int] = []

            async def tracking_process_page(page):
                process_page_calls.append(page.page_number)
                return await original_process_page(page)

            with (
                patch(
                    "apps.ocr.document_intelligence.pipeline.load_document",
                    return_value=fake_document,
                ),
                patch(
                    "apps.ocr.document_intelligence.pipeline.checkpoint_store.load_pages",
                    AsyncMock(return_value=checkpoints),
                ),
                patch(
                    "apps.ocr.document_intelligence.layout.LayoutDetector.detect_page",
                    _fake_detect_page,
                ),
                patch(
                    "apps.ocr.document_intelligence.elements.ElementProcessor._vlm_call",
                    _fake_vlm_call,
                ),
                patch.object(pipeline, "_process_page", tracking_process_page),
            ):
                result = await pipeline.process(
                    _blank_image(), "test.pdf", task_uid="task-1"
                )
        finally:
            pipeline.cleanup()

        assert "cached page one text" in result.markdown
        # Only page 2 should have actually been processed; page 1 came from
        # the checkpoint.
        assert process_page_calls == [2]


def _fake_document_with_pages(n_pages: int):
    from apps.ocr.document_intelligence.loader import Document, Page

    pages = [
        Page(
            id=f"p{i}",
            page_number=i,
            image=Image.new("RGB", (600, 800), "white"),
            width=600,
            height=800,
        )
        for i in range(1, n_pages + 1)
    ]
    return Document(
        id="doc1", filename="test.pdf", pages=pages, pages_count=n_pages, created_at="now"
    )


@pytest.mark.document_intelligence
class TestPageConcurrency:
    """
    Pages must no longer run strictly one-at-a-time.

    Layout detection stays serialized per document (the PaddleX predictor
    backing ``LayoutDetector`` mutates shared, unsynchronized instance state
    on every call -- see ``DocumentIntelligencePipeline._detect_layout``),
    but each page's VLM/HTTP element-extraction step must overlap with
    other pages' layout detection and VLM calls instead of the whole
    document running page-by-page.
    """

    async def test_pages_overlap_instead_of_running_strictly_sequentially(
        self,
    ) -> None:
        n_pages = 4
        delay = 0.15

        async def slow_vlm_call(
            _self,
            _crop,
            _system_prompt,
            _user_prompt,
            response_format=None,
            max_tokens: int = 1024,
        ) -> str:
            await asyncio.sleep(delay)
            return "some text"

        fake_document = _fake_document_with_pages(n_pages)

        pipeline = DocumentIntelligencePipeline(dpi=100)
        try:
            with (
                patch(
                    "apps.ocr.document_intelligence.pipeline.load_document",
                    return_value=fake_document,
                ),
                patch(
                    "apps.ocr.document_intelligence.layout.LayoutDetector.detect_page",
                    _fake_detect_page,
                ),
                patch(
                    "apps.ocr.document_intelligence.elements.ElementProcessor._vlm_call",
                    slow_vlm_call,
                ),
            ):
                start = time.monotonic()
                result = await pipeline.process(_blank_image(), "test.pdf")
                elapsed = time.monotonic() - start
        finally:
            pipeline.cleanup()

        assert len(result.stats.pages) == n_pages
        # Fully sequential (old behavior) would take at least n_pages * delay
        # (each page's VLM stage alone, ignoring layout time). Overlapping
        # pages should land well under half of that.
        serial_time = n_pages * delay
        assert elapsed < serial_time * 0.6, (
            f"elapsed={elapsed:.3f}s not much faster than fully-serial "
            f"{serial_time:.3f}s -- pages look strictly sequential"
        )

    async def test_page_order_preserved_regardless_of_completion_order(self) -> None:
        """
        Concurrent completion order must not scramble page order in the
        output -- the document AST/markdown must still read page 1..N in
        order even though pages may *finish* out of order (VLM delay makes
        later-submitted pages likely to finish before earlier ones whose
        layout-detection lock wait was longer).
        """
        n_pages = 5

        async def slow_vlm_call(
            _self,
            _crop,
            _system_prompt,
            _user_prompt,
            response_format=None,
            max_tokens: int = 1024,
        ) -> str:
            await asyncio.sleep(0.02)
            return "some text"

        fake_document = _fake_document_with_pages(n_pages)

        pipeline = DocumentIntelligencePipeline(dpi=100)
        try:
            with (
                patch(
                    "apps.ocr.document_intelligence.pipeline.load_document",
                    return_value=fake_document,
                ),
                patch(
                    "apps.ocr.document_intelligence.layout.LayoutDetector.detect_page",
                    _fake_detect_page,
                ),
                patch(
                    "apps.ocr.document_intelligence.elements.ElementProcessor._vlm_call",
                    slow_vlm_call,
                ),
            ):
                result = await pipeline.process(_blank_image(), "test.pdf")
        finally:
            pipeline.cleanup()

        assert [p.page_number for p in result.document_ast.pages] == list(
            range(1, n_pages + 1)
        )
        assert [p.page_number for p in result.stats.pages] == list(
            range(1, n_pages + 1)
        )

    async def test_progress_callback_stays_monotonic_and_uncounted_under_concurrency(
        self,
    ) -> None:
        """
        ``on_page_done`` must still be called exactly once per page, with a
        strictly increasing, never-repeating ``completed`` count -- even
        though pages may finish in a different order than they were
        submitted in under concurrency.
        """
        n_pages = 6
        delay = 0.02

        async def slow_vlm_call(
            _self,
            _crop,
            _system_prompt,
            _user_prompt,
            response_format=None,
            max_tokens: int = 1024,
        ) -> str:
            await asyncio.sleep(delay)
            return "some text"

        fake_document = _fake_document_with_pages(n_pages)
        progress_calls: list[tuple[int, int]] = []

        async def on_page_done(completed: int, total: int) -> None:
            progress_calls.append((completed, total))

        pipeline = DocumentIntelligencePipeline(dpi=100)
        try:
            with (
                patch(
                    "apps.ocr.document_intelligence.pipeline.load_document",
                    return_value=fake_document,
                ),
                patch(
                    "apps.ocr.document_intelligence.layout.LayoutDetector.detect_page",
                    _fake_detect_page,
                ),
                patch(
                    "apps.ocr.document_intelligence.elements.ElementProcessor._vlm_call",
                    slow_vlm_call,
                ),
            ):
                await pipeline.process(
                    _blank_image(), "test.pdf", on_page_done=on_page_done
                )
        finally:
            pipeline.cleanup()

        assert progress_calls == [(i, n_pages) for i in range(1, n_pages + 1)]

    async def test_checkpointed_page_skipped_while_others_run_concurrently(
        self,
    ) -> None:
        """Checkpoint-skip semantics must hold even with concurrent pages."""
        n_pages = 4
        fake_document = _fake_document_with_pages(n_pages)
        checkpoints = {
            1: {
                "page_ast": {
                    "page_number": 1,
                    "nodes": [
                        {
                            "type": "paragraph",
                            "text": "cached page one text",
                            "children": [],
                        }
                    ],
                },
                "page_stat": {
                    "page_number": 1,
                    "layout_time": 0.0,
                    "vlm_time": 0.0,
                },
            }
        }

        async def slow_vlm_call(
            _self,
            _crop,
            _system_prompt,
            _user_prompt,
            response_format=None,
            max_tokens: int = 1024,
        ) -> str:
            await asyncio.sleep(0.02)
            return "some text"

        pipeline = DocumentIntelligencePipeline(dpi=100)
        try:
            original_process_page = pipeline._process_page
            process_page_calls: list[int] = []

            async def tracking_process_page(page):
                process_page_calls.append(page.page_number)
                if page.page_number == 3:
                    error = RuntimeError("VLM exhausted retries")
                    raise error
                return await original_process_page(page)

            with (
                patch(
                    "apps.ocr.document_intelligence.pipeline.load_document",
                    return_value=fake_document,
                ),
                patch(
                    "apps.ocr.document_intelligence.pipeline.checkpoint_store.load_pages",
                    AsyncMock(return_value=checkpoints),
                ),
                patch(
                    "apps.ocr.document_intelligence.layout.LayoutDetector.detect_page",
                    _fake_detect_page,
                ),
                patch(
                    "apps.ocr.document_intelligence.elements.ElementProcessor._vlm_call",
                    slow_vlm_call,
                ),
                patch.object(pipeline, "_process_page", tracking_process_page),
            ):
                result = await pipeline.process(
                    _blank_image(), "test.pdf", task_uid="task-1"
                )
        finally:
            pipeline.cleanup()

        assert "cached page one text" in result.markdown
        # Page 1 came from the checkpoint and must not be reprocessed; pages
        # 2-4 (including the failing page 3) run concurrently.
        assert sorted(process_page_calls) == [2, 3, 4]
        failed = [p.page_number for p in result.stats.pages if p.failed]
        assert failed == [3]
        succeeded = {p.page_number for p in result.stats.pages if not p.failed}
        assert succeeded == {1, 2, 4}
        assert [p.page_number for p in result.stats.pages] == [1, 2, 3, 4]
