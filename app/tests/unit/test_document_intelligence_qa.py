"""
Unit tests for the post-generation DOCX QA gate (Phase 1 acceptance gate).

Each check is exercised both for the happy path (semantic renderer output
passes) and for a genuine failure it must catch — since the whole point of
this module is to fail loudly on regressions like "silently dropped a
block" or "text ended up in a text box", not just to pass on well-formed
input.
"""

from io import BytesIO

import pytest
from docx import Document as WordDocument

from apps.ocr.document_intelligence.ast import ASTNode, DocumentAST, PageAST
from apps.ocr.document_intelligence.layout import LayoutType
from apps.ocr.document_intelligence.qa import run_docx_qa
from apps.ocr.document_intelligence.renderers.docx import render_docx
from apps.ocr.document_intelligence.renderers.docx_absolute import render_docx_absolute


def _render(ast: DocumentAST) -> bytes:
    return render_docx(ast).getvalue()


@pytest.mark.document_intelligence
class TestSemanticOutputPassesEveryCheck:
    def test_mixed_document_passes_all_checks(self) -> None:
        pages = [
            PageAST(
                page_number=1,
                nodes=[
                    ASTNode(type=LayoutType.title, text="گزارش نمونه"),
                    ASTNode(
                        type=LayoutType.paragraph, text="این یک پاراگراف آزمایشی است."
                    ),
                    ASTNode(
                        type=LayoutType.table,
                        rows=[["A", "B"], ["1", "2"]],
                    ),
                    ASTNode(type=LayoutType.formula, latex=r"\frac{x^2}{y}"),
                ],
            )
        ]
        doc_ast = DocumentAST(pages=pages)

        report = run_docx_qa(doc_ast, _render(doc_ast))

        assert report.passed, report.failures()


@pytest.mark.document_intelligence
class TestNoTextBoxCheck:
    def test_semantic_output_has_no_text_boxes(self) -> None:
        node = ASTNode(type=LayoutType.paragraph, text="hello")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        report = run_docx_qa(doc_ast, _render(doc_ast))

        check = next(c for c in report.checks if c.name == "no_text_boxes")
        assert check.passed

    def test_visual_mode_textboxes_are_skipped_not_failed(self) -> None:
        node = ASTNode(type=LayoutType.paragraph, text="hello", bbox=(0, 0, 200, 50))
        doc_ast = DocumentAST(
            pages=[
                PageAST(
                    page_number=1,
                    nodes=[node],
                    page_width=1000,
                    page_height=1400,
                    page_dpi=200,
                )
            ]
        )

        report = run_docx_qa(
            doc_ast, render_docx_absolute(doc_ast).getvalue(), mode="visual"
        )

        check = next(c for c in report.checks if c.name == "no_text_boxes")
        assert check.passed
        assert "skipped" in check.detail


@pytest.mark.document_intelligence
class TestAllTextConsumedCheck:
    def test_catches_a_dropped_block(self) -> None:
        """
        Simulates the exact regression this check exists to catch: a
        node whose text never made it into the rendered document at all."""
        node = ASTNode(type=LayoutType.paragraph, text="این متن نباید گم شود")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        empty_docx = _render(DocumentAST(pages=[PageAST(page_number=1, nodes=[])]))
        report = run_docx_qa(doc_ast, empty_docx)

        check = next(c for c in report.checks if c.name == "all_ast_text_consumed")
        assert not check.passed
        assert not report.passed


@pytest.mark.document_intelligence
class TestReadingOrderCheck:
    def test_catches_out_of_order_paragraphs(self) -> None:
        doc_ast = DocumentAST(
            pages=[
                PageAST(
                    page_number=1,
                    nodes=[
                        ASTNode(type=LayoutType.paragraph, text="اول"),
                        ASTNode(type=LayoutType.paragraph, text="دوم"),
                    ],
                )
            ]
        )
        # Render normally, then swap the two body paragraphs' text to
        # simulate an out-of-order regression.
        doc = WordDocument(BytesIO(_render(doc_ast)))
        p0, p1 = doc.paragraphs[0], doc.paragraphs[1]
        p0.runs[0].text, p1.runs[0].text = p1.runs[0].text, p0.runs[0].text
        buf = BytesIO()
        doc.save(buf)

        report = run_docx_qa(doc_ast, buf.getvalue())

        check = next(c for c in report.checks if c.name == "reading_order_preserved")
        assert not check.passed


@pytest.mark.document_intelligence
class TestTablesAreRealCheck:
    def test_catches_missing_table(self) -> None:
        table_node = ASTNode(type=LayoutType.table, rows=[["A", "B"]])
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[table_node])])

        # Render a document with no table at all, but claim (via doc_ast)
        # that one was expected.
        blank_docx = _render(DocumentAST(pages=[PageAST(page_number=1, nodes=[])]))
        report = run_docx_qa(doc_ast, blank_docx)

        check = next(
            c for c in report.checks if c.name == "tables_are_real_word_tables"
        )
        assert not check.passed


@pytest.mark.document_intelligence
class TestFormulasAreRealOmmlCheck:
    def test_catches_missing_omml(self) -> None:
        formula_node = ASTNode(type=LayoutType.formula, latex=r"\frac{x^2}{y}")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[formula_node])])

        blank_docx = _render(DocumentAST(pages=[PageAST(page_number=1, nodes=[])]))
        report = run_docx_qa(doc_ast, blank_docx)

        check = next(c for c in report.checks if c.name == "formulas_are_real_omml")
        assert not check.passed


@pytest.mark.document_intelligence
class TestRtlBidiCheck:
    def test_catches_persian_paragraph_without_bidi(self) -> None:
        node = ASTNode(type=LayoutType.paragraph, text="سلام دنیا")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])
        doc = WordDocument(BytesIO(_render(doc_ast)))

        # Strip <w:bidi/> from the paragraph to simulate a regression.
        from docx.oxml.ns import qn

        p = next(p for p in doc.paragraphs if p.text == "سلام دنیا")
        ppr = p._p.find(qn("w:pPr"))
        bidi = ppr.find(qn("w:bidi"))
        ppr.remove(bidi)
        buf = BytesIO()
        doc.save(buf)

        report = run_docx_qa(doc_ast, buf.getvalue())

        check = next(c for c in report.checks if c.name == "rtl_bidi_correct")
        assert not check.passed


@pytest.mark.document_intelligence
class TestHeaderFooterPromotionCheck:
    def test_promoted_header_not_leaked_into_body(self) -> None:
        pages = [
            PageAST(
                page_number=n,
                nodes=[
                    ASTNode(type=LayoutType.header, text="Company Confidential"),
                    ASTNode(type=LayoutType.paragraph, text=f"body {n}"),
                ],
            )
            for n in (1, 2)
        ]
        doc_ast = DocumentAST(pages=pages)

        report = run_docx_qa(doc_ast, _render(doc_ast))

        check = next(
            c for c in report.checks if c.name == "header_footer_in_real_section"
        )
        assert check.passed

    def test_catches_header_leaking_into_body(self) -> None:
        pages = [
            PageAST(
                page_number=n,
                nodes=[ASTNode(type=LayoutType.header, text="Company Confidential")],
            )
            for n in (1, 2)
        ]
        doc_ast = DocumentAST(pages=pages)
        doc = WordDocument(BytesIO(_render(doc_ast)))
        doc.add_paragraph("Company Confidential")  # simulate a leak into the body
        buf = BytesIO()
        doc.save(buf)

        report = run_docx_qa(doc_ast, buf.getvalue())

        check = next(
            c for c in report.checks if c.name == "header_footer_in_real_section"
        )
        assert not check.passed
        assert "leaked" in check.detail


@pytest.mark.document_intelligence
class TestPageNumberFieldCheck:
    def test_verified_page_number_sequence_passes(self) -> None:
        pages = [
            PageAST(
                page_number=n, nodes=[ASTNode(type=LayoutType.page_number, text=str(n))]
            )
            for n in (1, 2, 3)
        ]
        doc_ast = DocumentAST(pages=pages)

        report = run_docx_qa(doc_ast, _render(doc_ast))

        check = next(
            c for c in report.checks if c.name == "page_number_uses_real_field"
        )
        assert check.passed
        assert "skipped" not in check.detail

    def test_no_page_number_sequence_is_skipped_not_failed(self) -> None:
        node = ASTNode(type=LayoutType.paragraph, text="hello")
        doc_ast = DocumentAST(pages=[PageAST(page_number=1, nodes=[node])])

        report = run_docx_qa(doc_ast, _render(doc_ast))

        check = next(
            c for c in report.checks if c.name == "page_number_uses_real_field"
        )
        assert check.passed
        assert "skipped" in check.detail

    def test_catches_missing_page_field_when_one_was_expected(self) -> None:
        """
        Simulates a renderer regression: a verified page-number sequence
        exists in the AST, but the output never actually got a PAGE field."""
        pages = [
            PageAST(
                page_number=n, nodes=[ASTNode(type=LayoutType.page_number, text=str(n))]
            )
            for n in (1, 2, 3)
        ]
        doc_ast = DocumentAST(pages=pages)

        blank_docx = _render(
            DocumentAST(pages=[PageAST(page_number=n, nodes=[]) for n in (1, 2, 3)])
        )
        report = run_docx_qa(doc_ast, blank_docx)

        check = next(
            c for c in report.checks if c.name == "page_number_uses_real_field"
        )
        assert not check.passed
