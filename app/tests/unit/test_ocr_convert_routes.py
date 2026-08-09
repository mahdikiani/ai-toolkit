"""
Unit tests for Markdown -> DOCX conversion route handlers (direct call,
no HTTP/auth layer -- see tests/integration/test_convert_routes.py for the
auth-gating/wiring checks)."""

from io import BytesIO

import pytest
from docx import Document as WordDocument
from usso import UserData

from apps.ocr.convert_routes import (
    MarkdownToDocxRequest,
    markdown_to_docx,
    markdown_to_docx_upload,
    markdown_to_pdf,
    markdown_to_pdf_upload,
)

_FAKE_USER = UserData(sub="test-user")


class _FakeUploadFile:
    """Minimal stand-in for FastAPI's UploadFile in a direct unit-test call."""

    def __init__(self, content: bytes, filename: str) -> None:
        self._content = content
        self.filename = filename

    async def read(self) -> bytes:
        return self._content


async def _collect(streaming_response) -> bytes:
    chunks = [chunk async for chunk in streaming_response.body_iterator]
    return b"".join(
        chunk if isinstance(chunk, bytes) else chunk.encode("utf-8") for chunk in chunks
    )


@pytest.mark.unit
class TestMarkdownToDocxTextEndpoint:
    async def test_produces_a_real_editable_docx(self) -> None:
        data = MarkdownToDocxRequest(
            markdown="# عنوان\n\nاین یک **پاراگراف** آزمایشی است.", title="گزارش"
        )

        response = await markdown_to_docx(data, user=_FAKE_USER)
        body = await _collect(response)
        doc = WordDocument(BytesIO(body))

        assert doc.core_properties.title == "گزارش"
        assert any("عنوان" in p.text for p in doc.paragraphs)
        bold_runs = [r for p in doc.paragraphs for r in p.runs if r.bold]
        assert bold_runs

    def test_content_disposition_and_media_type(self) -> None:
        import asyncio

        data = MarkdownToDocxRequest(markdown="hello", title="")
        response = asyncio.run(markdown_to_docx(data, user=_FAKE_USER))

        assert response.media_type == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        assert 'filename="document.docx"' in response.headers["Content-Disposition"]


@pytest.mark.unit
class TestMarkdownToDocxUploadEndpoint:
    async def test_uploaded_file_content_is_converted(self) -> None:
        upload = _FakeUploadFile(
            "# سلام\n\nمتن آزمایشی از فایل.".encode(), "notes.md"
        )

        response = await markdown_to_docx_upload(file=upload, title="", user=_FAKE_USER)
        body = await _collect(response)
        doc = WordDocument(BytesIO(body))

        assert any("سلام" in p.text for p in doc.paragraphs)

    async def test_title_defaults_to_filename_stem_when_not_given(self) -> None:
        upload = _FakeUploadFile(b"content", "meeting-notes.md")

        response = await markdown_to_docx_upload(file=upload, title="", user=_FAKE_USER)

        assert 'filename="meeting-notes.docx"' in response.headers["Content-Disposition"]
        body = await _collect(response)
        doc = WordDocument(BytesIO(body))
        assert doc.core_properties.title == "meeting-notes"

    async def test_explicit_title_overrides_filename_stem(self) -> None:
        upload = _FakeUploadFile(b"content", "raw_upload_123.md")

        response = await markdown_to_docx_upload(
            file=upload, title="عنوان دلخواه", user=_FAKE_USER
        )

        # Non-ASCII file names can't appear raw in an HTTP header (Latin-1
        # only) -- RFC 5987/6266 extended filename*, percent-encoded.
        disposition = response.headers["Content-Disposition"]
        assert "filename*=UTF-8''" in disposition
        from urllib.parse import unquote

        encoded = disposition.split("filename*=UTF-8''", 1)[1]
        assert unquote(encoded) == "عنوان دلخواه.docx"

    async def test_non_ascii_title_does_not_crash_the_response(self) -> None:
        """
        Regression: a Persian filename/title used to raise
        UnicodeEncodeError deep inside Starlette's header encoding,
        turning a valid conversion into a 500."""
        upload = _FakeUploadFile(b"content", "یادداشت.md")

        response = await markdown_to_docx_upload(file=upload, title="", user=_FAKE_USER)
        body = await _collect(response)  # must not raise

        assert len(body) > 0


@pytest.mark.unit
class TestMarkdownToPdfEndpoints:
    async def test_json_endpoint_returns_pdf_with_ascii_filename(self) -> None:
        data = MarkdownToDocxRequest(markdown="# Report\n\nPDF body", title="Report")

        response = await markdown_to_pdf(data, user=_FAKE_USER)
        body = await _collect(response)

        assert response.status_code == 200
        assert response.media_type == "application/pdf"
        assert 'filename="document.pdf"' in response.headers["Content-Disposition"]
        assert body.startswith(b"%PDF-")

    async def test_upload_endpoint_returns_pdf_with_persian_filename(self) -> None:
        from urllib.parse import unquote

        upload = _FakeUploadFile("# عنوان\n\nمتن فارسی".encode(), "notes.md")

        response = await markdown_to_pdf_upload(
            file=upload,
            title="گزارش فارسی",
            user=_FAKE_USER,
        )
        body = await _collect(response)
        disposition = response.headers["Content-Disposition"]

        assert response.status_code == 200
        assert response.media_type == "application/pdf"
        assert "filename*=UTF-8''" in disposition
        encoded = disposition.split("filename*=UTF-8''", 1)[1]
        assert unquote(encoded) == "گزارش فارسی.pdf"
        assert body.startswith(b"%PDF-")
